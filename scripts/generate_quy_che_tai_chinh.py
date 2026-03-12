#!/usr/bin/env python3
"""Generate Quy Chế Tài Chính Nội Bộ for OSR"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_document():
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("CÔNG TY OS RESEARCH (OSR)")
    run.bold = True
    run.font.size = Pt(14)

    header2 = doc.add_paragraph()
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = header2.add_run("───────────────")

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("\nQUY CHẾ TÀI CHÍNH NỘI BỘ")
    run_title.bold = True
    run_title.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run(f"(Ban hành kèm theo Quyết định số: 01/2026/QĐ-OSR ngày {datetime.now().strftime('%d/%m/%Y')})")
    run_sub.italic = True
    run_sub.font.size = Pt(12)

    doc.add_paragraph()

    # CHƯƠNG I
    add_chapter(doc, "CHƯƠNG I", "QUY ĐỊNH CHUNG")

    add_article(doc, "Điều 1. Mục đích", """
Quy chế này quy định các nguyên tắc, quy trình quản lý tài chính nội bộ của Công ty OS Research (sau đây gọi là "Công ty" hoặc "OSR"), nhằm:
1. Đảm bảo sử dụng nguồn lực tài chính hiệu quả, minh bạch.
2. Tạo cơ sở pháp lý cho các hoạt động thu chi trong Công ty.
3. Phân định rõ trách nhiệm của các cá nhân, bộ phận liên quan.
""")

    add_article(doc, "Điều 2. Phạm vi và đối tượng áp dụng", """
1. Phạm vi: Tất cả hoạt động tài chính của Công ty.
2. Đối tượng: Toàn bộ thành viên, cộng tác viên và các bên liên quan đến hoạt động tài chính của Công ty.
""")

    add_article(doc, "Điều 3. Nguyên tắc quản lý tài chính", """
1. Minh bạch: Mọi khoản thu chi phải được ghi chép đầy đủ, rõ ràng.
2. Tiết kiệm: Sử dụng nguồn lực tài chính hiệu quả, tránh lãng phí.
3. Đúng mục đích: Chi tiêu phải phục vụ hoạt động kinh doanh của Công ty.
4. Có chứng từ: Mọi khoản chi phải có hóa đơn, chứng từ hợp lệ.
5. Được phê duyệt: Các khoản chi phải được người có thẩm quyền phê duyệt trước khi thực hiện.
""")

    # CHƯƠNG II
    add_chapter(doc, "CHƯƠNG II", "QUẢN LÝ THU CHI")

    add_article(doc, "Điều 4. Quản lý nguồn thu", """
1. Các nguồn thu của Công ty bao gồm:
   a) Doanh thu từ sản phẩm, dịch vụ.
   b) Vốn góp của thành viên.
   c) Vốn đầu tư từ bên ngoài.
   d) Các nguồn thu hợp pháp khác.

2. Tất cả nguồn thu phải được nộp vào tài khoản chính thức của Công ty trong vòng 24 giờ kể từ khi nhận được.

3. Nghiêm cấm việc sử dụng nguồn thu cho mục đích cá nhân trước khi nộp về Công ty.
""")

    add_article(doc, "Điều 5. Phân loại chi phí", """
1. Chi phí hoạt động thường xuyên:
   a) Lương và phụ cấp cho nhân sự.
   b) Chi phí thuê văn phòng, điện, nước, internet.
   c) Chi phí công cụ, phần mềm phục vụ công việc.
   d) Chi phí marketing, quảng cáo.

2. Chi phí dự án/thí nghiệm:
   a) Chi phí nghiên cứu và phát triển sản phẩm.
   b) Chi phí chạy thí nghiệm (ads, landing page, prototype...).
   c) Chi phí thuê ngoài (freelancer, agency...).

3. Chi phí đầu tư:
   a) Mua sắm tài sản, thiết bị.
   b) Đầu tư vào các dự án nội bộ (Internal Projects).
""")

    add_article(doc, "Điều 6. Thẩm quyền phê duyệt chi", """
1. Giám đốc/Founder phê duyệt:
   a) Tất cả các khoản chi trên 10.000.000 VNĐ (mười triệu đồng).
   b) Chi phí đầu tư, mua sắm tài sản.
   c) Hợp đồng với bên thứ ba.

2. Trưởng bộ phận/Project Lead được phê duyệt:
   a) Các khoản chi dưới 10.000.000 VNĐ thuộc ngân sách được giao.
   b) Chi phí vận hành thường xuyên trong phạm vi dự toán.

3. Cá nhân tự chi và hoàn ứng:
   a) Các khoản chi dưới 2.000.000 VNĐ (hai triệu đồng) phục vụ công việc.
   b) Phải có hóa đơn/chứng từ và được phê duyệt hoàn ứng trong vòng 7 ngày.
""")

    add_article(doc, "Điều 7. Quy trình thanh toán", """
1. Đề xuất chi:
   a) Người đề xuất điền Phiếu đề xuất chi (mẫu kèm theo).
   b) Ghi rõ: mục đích, số tiền, thời gian cần thanh toán.

2. Phê duyệt:
   a) Chuyển đề xuất cho người có thẩm quyền theo Điều 6.
   b) Thời gian phê duyệt tối đa: 2 ngày làm việc.

3. Thanh toán:
   a) Ưu tiên chuyển khoản ngân hàng.
   b) Thanh toán tiền mặt chỉ áp dụng cho khoản chi dưới 2.000.000 VNĐ.

4. Lưu chứng từ:
   a) Scan/chụp hóa đơn, chứng từ gửi về email/folder chung.
   b) Lưu bản gốc (nếu có) tại bộ phận kế toán.
""")

    # CHƯƠNG III
    add_chapter(doc, "CHƯƠNG III", "QUẢN LÝ NGÂN SÁCH")

    add_article(doc, "Điều 8. Lập ngân sách", """
1. Ngân sách năm:
   a) Lập vào tháng 12 hàng năm cho năm tiếp theo.
   b) Bao gồm: dự toán thu, dự toán chi theo từng hạng mục.
   c) Giám đốc phê duyệt ngân sách năm.

2. Ngân sách dự án/thí nghiệm:
   a) Lập trước khi bắt đầu dự án.
   b) Project Lead chịu trách nhiệm quản lý ngân sách được giao.
   c) Báo cáo tiến độ sử dụng ngân sách hàng tuần.
""")

    add_article(doc, "Điều 9. Kiểm soát ngân sách", """
1. Không được chi vượt ngân sách đã được phê duyệt.
2. Nếu cần chi vượt, phải:
   a) Lập đề xuất điều chỉnh ngân sách.
   b) Được Giám đốc phê duyệt trước khi thực hiện.
3. Cuối mỗi quý, rà soát và điều chỉnh ngân sách nếu cần.
""")

    # CHƯƠNG IV
    add_chapter(doc, "CHƯƠNG IV", "QUẢN LÝ TÀI SẢN")

    add_article(doc, "Điều 10. Tài sản của Công ty", """
1. Tài sản bao gồm:
   a) Tiền mặt, tiền gửi ngân hàng.
   b) Thiết bị, công cụ (máy tính, camera, thiết bị văn phòng...).
   c) Tài sản trí tuệ (thương hiệu, phần mềm, nội dung...).

2. Quy định sử dụng:
   a) Tài sản Công ty chỉ được sử dụng cho mục đích công việc.
   b) Người được giao tài sản chịu trách nhiệm bảo quản.
   c) Mất mát, hư hỏng do lỗi cá nhân phải bồi thường.
""")

    # CHƯƠNG V
    add_chapter(doc, "CHƯƠNG V", "BÁO CÁO VÀ KIỂM TRA")

    add_article(doc, "Điều 11. Chế độ báo cáo", """
1. Báo cáo hàng tuần:
   a) Tổng hợp thu chi trong tuần.
   b) Cập nhật tình hình ngân sách dự án.

2. Báo cáo hàng tháng:
   a) Báo cáo thu chi chi tiết.
   b) So sánh thực tế với ngân sách.
   c) Gửi cho Giám đốc trước ngày 5 tháng sau.

3. Báo cáo hàng quý:
   a) Tổng kết tài chính quý.
   b) Đề xuất điều chỉnh ngân sách (nếu có).
""")

    add_article(doc, "Điều 12. Kiểm tra tài chính", """
1. Giám đốc có quyền kiểm tra tài chính bất kỳ lúc nào.
2. Tất cả thành viên có nghĩa vụ hợp tác, cung cấp thông tin khi được yêu cầu.
3. Phát hiện sai phạm sẽ xử lý theo quy định tại Điều 13.
""")

    # CHƯƠNG VI
    add_chapter(doc, "CHƯƠNG VI", "XỬ LÝ VI PHẠM")

    add_article(doc, "Điều 13. Xử lý vi phạm", """
1. Vi phạm nhẹ (lần đầu, không gây thiệt hại):
   a) Nhắc nhở, cảnh cáo bằng văn bản.

2. Vi phạm nghiêm trọng:
   a) Chi sai mục đích, không có phê duyệt: Hoàn trả 100% số tiền + cảnh cáo.
   b) Làm mất tài sản: Bồi thường theo giá trị tài sản.
   c) Gian lận, biển thủ: Chấm dứt hợp tác + bồi thường + xử lý theo pháp luật.
""")

    # CHƯƠNG VII
    add_chapter(doc, "CHƯƠNG VII", "ĐIỀU KHOẢN THI HÀNH")

    add_article(doc, "Điều 14. Hiệu lực", """
1. Quy chế này có hiệu lực từ ngày ký ban hành.
2. Mọi quy định trước đây trái với Quy chế này đều bị bãi bỏ.
3. Việc sửa đổi, bổ sung Quy chế do Giám đốc quyết định.
""")

    add_article(doc, "Điều 15. Trách nhiệm thi hành", """
1. Tất cả thành viên Công ty có trách nhiệm nghiên cứu, thực hiện đúng Quy chế này.
2. Bộ phận kế toán/tài chính chịu trách nhiệm hướng dẫn, giám sát việc thực hiện.
3. Mọi thắc mắc liên hệ trực tiếp với Giám đốc để được giải đáp.
""")

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature section
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = True

    cell1 = sig_table.rows[0].cells[0]
    cell2 = sig_table.rows[0].cells[1]

    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"TP. Hồ Chí Minh, ngày {datetime.now().strftime('%d')} tháng {datetime.now().strftime('%m')} năm {datetime.now().strftime('%Y')}\n")
    run2.font.size = Pt(13)
    run2b = p2.add_run("GIÁM ĐỐC\n\n\n\n\n")
    run2b.bold = True
    run2b.font.size = Pt(13)
    run2c = p2.add_run("_______________________")
    run2c.font.size = Pt(13)

    return doc

def add_chapter(doc, chapter_num, chapter_title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n{chapter_num}")
    run.bold = True
    run.font.size = Pt(14)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(chapter_title)
    run2.bold = True
    run2.font.size = Pt(14)

def add_article(doc, title, content):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)

    # Add content
    for line in content.strip().split('\n'):
        p2 = doc.add_paragraph(line.strip())
        p2.paragraph_format.left_indent = Inches(0.25)
        p2.paragraph_format.space_after = Pt(3)

if __name__ == "__main__":
    doc = create_document()
    output_path = "/Users/Vincent/OSR-company/os-research-engine/docs/Quy_che_tai_chinh_noi_bo_OSR.docx"

    # Ensure directory exists
    import os
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc.save(output_path)
    print(f"Document saved to: {output_path}")
