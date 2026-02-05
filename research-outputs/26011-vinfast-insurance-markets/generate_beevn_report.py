#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
BeeVN Insurance Market Report Generator
Generates Vietnamese-language professional report for Vinfast insurance markets
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def add_formatted_heading(doc, text, level=1):
    """Add a formatted heading to the document"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph"""
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
    return para

def add_table(doc, data, headers=None):
    """Add a formatted table"""
    rows = len(data) + (1 if headers else 0)
    cols = len(data[0]) if data else 0

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'

    # Add headers
    if headers:
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)

    # Add data
    start_row = 1 if headers else 0
    for i, row_data in enumerate(data):
        row_cells = table.rows[start_row + i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = str(cell_data)
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)

    return table

def create_cover_page(doc):
    """Create the cover page"""
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('BÁO CÁO PHÂN TÍCH\nTHỊ TRƯỜNG BẢO HIỂM PHI NHÂN THỌ')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_run.font.name = 'Arial'

    doc.add_paragraph()

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Congo (DRC), Kazakhstan, Philippines')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
    subtitle_run.font.name = 'Arial'

    doc.add_paragraph()
    doc.add_paragraph()

    # Client info
    client = doc.add_paragraph()
    client.alignment = WD_ALIGN_PARAGRAPH.CENTER
    client_run = client.add_run('Khách hàng: BeeVN\nĐối tác Bảo hiểm Ủy quyền của Vinfast')
    client_run.font.size = Pt(14)
    client_run.font.name = 'Arial'

    doc.add_paragraph()

    # Date
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date.add_run('04 tháng 02 năm 2026')
    date_run.font.size = Pt(12)
    date_run.font.name = 'Arial'

    doc.add_page_break()

def create_executive_summary(doc):
    """Create executive summary section"""
    add_formatted_heading(doc, 'TÓM TẮT ĐIỀU HÀNH', level=1)

    para = doc.add_paragraph()
    para.add_run('Báo cáo này phân tích thị trường bảo hiểm phi nhân thọ tại ba thị trường trọng điểm cho sự mở rộng của Vinfast: Cộng Hòa Dân Chủ Congo (DRC), Kazakhstan, và Philippines. Dưới đây là những phát hiện chính và khuyến nghị chiến lược cho BeeVN.').font.name = 'Arial'

    doc.add_paragraph()
    add_formatted_heading(doc, 'So Sánh Nhanh Ba Thị Trường', level=2)

    comparison_data = [
        ['DRC (Congo)', 'ARCA', 'KHÔNG phải thành viên CIMA', 'Thành lập công ty địa phương bắt buộc', 'Chưa xác định', 'USD 350M (2025)'],
        ['Kazakhstan', 'ARDFM (NBK)', 'Thành viên EAEU', 'Chi nhánh hoặc công ty con', 'Khác nhau theo hình thức', 'KZT 674.5 tỷ (2025)'],
        ['Philippines', 'Insurance Commission', 'Độc lập', '100% sở hữu nước ngoài được phép', 'PHP 20M (~$350K)', 'USD 2.11 tỷ (2024)']
    ]

    headers = ['Thị trường', 'Cơ quan quản lý', 'Khung khu vực', 'Yêu cầu môi giới nước ngoài', 'Vốn tối thiểu', 'Quy mô thị trường']
    add_table(doc, comparison_data, headers)

    doc.add_paragraph()
    add_formatted_heading(doc, 'Những Phát Hiện Chính', level=2)

    findings = [
        ('DRC - Thị trường đang phát triển nhanh:', 'Thị trường bảo hiểm phi nhân thọ tăng trưởng 5 lần từ $67M lên $350M sau khi tự do hóa năm 2018. YÊU CẦU QUAN TRỌNG: Bảo hiểm hàng hóa nhập khẩu phải sử dụng công ty bảo hiểm địa phương DRC.'),
        ('Kazakhstan - Môi trường quy định mạnh:', 'Thị trường bảo hiểm tăng trưởng 24% năm 2025. Hơn 19,000 xe điện đã đăng ký (H1 2025) với các ưu đãi của chính phủ. Thị trường bảo hiểm EV chưa phát triển - cơ hội lớn cho chuyên môn hóa.'),
        ('Philippines - Hoạt động Vinfast đang diễn ra:', 'Vinfast đã mở các đại lý từ tháng 7/2024 với quan hệ đối tác BDO. Luật EVIDA cung cấp ưu đãi cho EV, yêu cầu 5% đội xe của công ty phải là xe điện. Doanh số EV tăng 80% năm 2024.')
    ]

    for title, content in findings:
        para = doc.add_paragraph()
        para.add_run(title).bold = True
        para.add_run(f' {content}')
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.add_page_break()

def create_drc_section(doc):
    """Create DRC section"""
    add_formatted_heading(doc, 'PHẦN 1: CỘNG HÒA DÂN CHỦ CONGO (DRC)', level=1)

    # Regulatory framework
    add_formatted_heading(doc, '1.1 Khung Pháp Lý', level=2)
    para = doc.add_paragraph()
    para.add_run('Cơ quan quản lý: ').bold = True
    para.add_run('ARCA (Autorité de Régulation et de Contrôle des Assurances)')
    for run in para.runs:
        run.font.name = 'Arial'

    para = doc.add_paragraph()
    para.add_run('Luật chính: ').bold = True
    para.add_run('Luật Bảo hiểm số 15/005 ngày 17 tháng 3 năm 2015')
    for run in para.runs:
        run.font.name = 'Arial'

    para = doc.add_paragraph()
    para.add_run('Tình trạng CIMA: ').bold = True
    para.add_run('DRC KHÔNG phải là thành viên CIMA - hoạt động theo hệ thống quản lý độc lập')
    for run in para.runs:
        run.font.name = 'Arial'

    doc.add_paragraph()

    # Licensing requirements
    add_formatted_heading(doc, '1.2 Yêu Cầu Cấp Phép Môi Giới', level=2)
    para = doc.add_paragraph()
    para.add_run('Cấu trúc pháp lý bắt buộc:').bold = True
    for run in para.runs:
        run.font.name = 'Arial'

    requirements = [
        'Phải thành lập công ty theo luật Congo dưới hình thức société anonyme (SA)',
        'Bổ nhiệm ban quản lý có trình độ (bằng đại học + 10 năm kinh nghiệm)',
        'Xin giấy phép ARCA với tư cách trung gian bảo hiểm',
        'Thiết lập văn phòng thực tế tại DRC',
        'Chưa rõ: Yêu cầu đối tác địa phương, vốn tối thiểu cụ thể'
    ]

    for req in requirements:
        para = doc.add_paragraph(req, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.add_paragraph()

    # Compulsory insurance
    add_formatted_heading(doc, '1.3 Bảo Hiểm Bắt Buộc', level=2)

    compulsory_data = [
        ['Bảo hiểm xe cơ giới', 'Ordinance 73-013 (1973)', 'Bảo hiểm trách nhiệm dân sự bên thứ ba', 'Tất cả chủ xe', 'Hàng năm'],
        ['Bồi thường người lao động', 'Chương trình nhà nước INSS', '1.5-3% lương', 'Tất cả nhà tuyển dụng', 'Hàng tháng'],
        ['Bảo hiểm xây dựng', 'Luật 74-007 (1974)', 'Trách nhiệm 10 năm (decennial)', 'Tất cả dự án xây dựng', 'Theo dự án'],
        ['Bảo hiểm hàng nhập khẩu', 'Quy định DRC', 'Bảo hiểm hàng hóa', 'Tất cả nhà nhập khẩu', 'Theo lô hàng']
    ]

    headers = ['Loại bảo hiểm', 'Cơ sở pháp lý', 'Phạm vi bảo hiểm', 'Đối tượng bắt buộc', 'Thời hạn']
    add_table(doc, compulsory_data, headers)

    doc.add_paragraph()

    # Critical requirement box
    para = doc.add_paragraph()
    para.add_run('⚠ YÊU CẦU QUAN TRỌNG: ').bold = True
    para.add_run('Bảo hiểm hàng hóa nhập khẩu PHẢI được mua từ công ty bảo hiểm có trụ sở tại DRC. Đây là yêu cầu duy nhất và bắt buộc đối với hoạt động nhập khẩu của Vinfast.')
    for run in para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.right_indent = Inches(0.5)

    doc.add_paragraph()

    # Market players
    add_formatted_heading(doc, '1.4 Thị Trường và Các Công Ty Bảo Hiểm', level=2)

    para = doc.add_paragraph()
    para.add_run('Tăng trưởng thị trường: ').bold = True
    para.add_run('Phí bảo hiểm phi nhân thọ tăng từ $67M (trước 2018) lên $350M (2025) - tăng 5 lần sau tự do hóa.')
    for run in para.runs:
        run.font.name = 'Arial'

    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run('Các công ty bảo hiểm chính:').bold = True
    for run in para.runs:
        run.font.name = 'Arial'

    companies = [
        'SFA (Société Financière d\'Assurance) - Công ty số 2 thị trường, 25% doanh thu từ ngành khai thác',
        'Rawsur DRC và Rawsur Life - Công ty dẫn đầu thị trường',
        'SONAS - Công ty nhà nước (độc quyền trước đây 1968-2018)',
        'Activa Assurance DRC - Thuộc Activa Group (Cameroon)',
        'Mayfair Insurance DRC - Hoạt động tích cực'
    ]

    for company in companies:
        para = doc.add_paragraph(company, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.add_paragraph()

    # Recommendations for DRC
    add_formatted_heading(doc, '1.5 Khuyến Nghị cho Vinfast/BeeVN tại DRC', level=2)

    recommendations = [
        'Liên hệ trực tiếp với ARCA để làm rõ yêu cầu cấp phép môi giới và vốn tối thiểu',
        'Thuê tư vấn pháp lý Congo để bắt đầu quy trình thành lập société anonyme',
        'Thiết lập quan hệ với các công ty bảo hiểm chính (SFA, Rawsur, SONAS)',
        'Tuân thủ yêu cầu bảo hiểm hàng nhập khẩu - sử dụng công ty bảo hiểm DRC cho tất cả lô hàng',
        'Cân nhắc bảo hiểm rủi ro chính trị do tình hình bất ổn khu vực (phí tăng 5-10 lần năm 2025)'
    ]

    for i, rec in enumerate(recommendations, 1):
        para = doc.add_paragraph(f'{i}. {rec}')
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.add_page_break()

def create_kazakhstan_section(doc):
    """Create Kazakhstan section"""
    add_formatted_heading(doc, 'PHẦN 2: KAZAKHSTAN', level=1)

    # Regulatory framework
    add_formatted_heading(doc, '2.1 Khung Pháp Lý', level=2)
    para = doc.add_paragraph()
    para.add_run('Cơ quan quản lý: ').bold = True
    para.add_run('ARDFM (Agency for Regulation and Development of the Financial Market) thuộc Ngân hàng Quốc gia Kazakhstan (NBK)')
    for run in para.runs:
        run.font.name = 'Arial'

    para = doc.add_paragraph()
    para.add_run('Luật chính: ').bold = True
    para.add_run('Luật "Về Hoạt động Bảo hiểm" số 126-II (2000), sửa đổi tháng 1/2023')
    for run in para.runs:
        run.font.name = 'Arial'

    para = doc.add_paragraph()
    para.add_run('Thành viên EAEU: ').bold = True
    para.add_run('Kazakhstan là thành viên Liên minh Kinh tế Á-Âu (cùng với Armenia, Belarus, Kyrgyzstan, Nga), nhưng vẫn yêu cầu cấp phép quốc gia riêng')
    for run in para.runs:
        run.font.name = 'Arial'

    doc.add_paragraph()

    # Licensing requirements
    add_formatted_heading(doc, '2.2 Yêu Cầu Cấp Phép Môi Giới Nước Ngoài', level=2)

    para = doc.add_paragraph()
    para.add_run('Môi giới bảo hiểm nước ngoài có thể mở chi nhánh tại Kazakhstan thông qua cổng thông tin điện tử: www.elicense.kz').font.name = 'Arial'

    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run('Hồ sơ yêu cầu:').bold = True
    for run in para.runs:
        run.font.name = 'Arial'

    requirements = [
        'Xác nhận bằng văn bản từ cơ quan có thẩm quyền của nước sở tại về hồ sơ sạch của người sáng lập',
        'Bản sao giấy phép môi giới bảo hiểm hợp lệ từ nước sở tại',
        'Hồ sơ đăng ký chi nhánh theo Điều 6-2 của Luật Kazakhstan',
        'Xác nhận thanh toán phí cấp phép'
    ]

    for req in requirements:
        para = doc.add_paragraph(req, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.add_paragraph()

    # Compulsory insurance
    add_formatted_heading(doc, '2.3 Bảo Hiểm Bắt Buộc', level=2)

    compulsory_data = [
        ['OSAGO (Xe cơ giới)', 'Luật 446-II (2003)', 'Trách nhiệm dân sự bên thứ ba', 'Tất cả chủ xe', '12 tháng'],
        ['Tai nạn nghề nghiệp', 'Luật 30 (2005)', '30-100% mất năng lực', 'Tất cả nhà tuyển dụng', 'Theo lương'],
        ['Trách nhiệm môi trường', 'Bộ luật Môi trường', 'Thiệt hại môi trường', 'Cơ sở Loại I (≥1,000 tấn khí thải)', '12 tháng']
    ]

    headers = ['Loại bảo hiểm', 'Cơ sở pháp lý', 'Phạm vi bảo hiểm', 'Đối tượng bắt buộc', 'Thời hạn']
    add_table(doc, compulsory_data, headers)

    doc.add_paragraph()

    # EV market
    add_formatted_heading(doc, '2.4 Thị Trường EV và Cơ Hội', level=2)

    para = doc.add_paragraph()
    para.add_run('Quy mô thị trường EV: ').bold = True
    para.add_run('Hơn 19,000 xe điện đã đăng ký (H1 2025), tăng trưởng nhanh')
    for run in para.runs:
        run.font.name = 'Arial'

    para = doc.add_paragraph()
    para.add_run('Hỗ trợ của chính phủ: ').bold = True
    para.add_run('Miễn thuế giao thông đến 31/12/2025, trợ cấp đầu tư 30% cho dự án cọc sạc, đảm bảo lợi nhuận tối thiểu 8%')
    for run in para.runs:
        run.font.name = 'Arial'

    para = doc.add_paragraph()
    para.add_run('Dự báo 2030: ').bold = True
    para.add_run('78,000-120,000 xe điện (12% đội xe quốc gia), cần 4,000-8,000 điểm sạc')
    for run in para.runs:
        run.font.name = 'Arial'

    doc.add_paragraph()

    # Major insurers
    add_formatted_heading(doc, '2.5 Các Công Ty Bảo Hiểm Lớn', level=2)

    insurers_data = [
        ['IC Centras Insurance JSC', '$67.8M', 'Công ty dẫn đầu thị trường'],
        ['JSC Eurasia Insurance', '$46.7M', 'Công ty phi nhân thọ mạnh'],
        ['Halyk Insurance Company JSC', '$31.1M', 'Thuộc tập đoàn Ngân hàng Halyk']
    ]

    headers = ['Công ty', 'Doanh thu (2025)', 'Ghi chú']
    add_table(doc, insurers_data, headers)

    doc.add_paragraph()

    # Recommendations
    add_formatted_heading(doc, '2.6 Khuyến Nghị cho Vinfast/BeeVN tại Kazakhstan', level=2)

    recommendations = [
        'Xin xác nhận từ cơ quan quản lý bảo hiểm Việt Nam về giấy phép hợp lệ và hồ sơ sạch',
        'Bắt đầu đăng ký trên cổng thông tin e-licensing tại www.elicense.kz',
        'Thuê tư vấn pháp lý địa phương chuyên về luật bảo hiểm',
        'Phỏng vấn các công ty bảo hiểm chính (IC Centras, Eurasia, Halyk) về sản phẩm bảo hiểm EV',
        'Định vị là chuyên gia bảo hiểm EV - thị trường chưa phát triển, cơ hội lớn cho chuyên môn hóa',
        'Mua bảo hiểm trách nhiệm nghề nghiệp (yêu cầu từ tháng 1/2025)'
    ]

    for i, rec in enumerate(recommendations, 1):
        para = doc.add_paragraph(f'{i}. {rec}')
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.add_page_break()

def create_philippines_section(doc):
    """Create Philippines section"""
    add_formatted_heading(doc, 'PHẦN 3: PHILIPPINES', level=1)

    # Regulatory framework
    add_formatted_heading(doc, '3.1 Khung Pháp Lý', level=2)
    para = doc.add_paragraph()
    para.add_run('Cơ quan quản lý: ').bold = True
    para.add_run('Insurance Commission (IC) thuộc Bộ Tài chính')
    for run in para.runs:
        run.font.name = 'Arial'

    para = doc.add_paragraph()
    para.add_run('Luật chính: ').bold = True
    para.add_run('Bộ luật Bảo hiểm (PD 612, sửa đổi bởi RA 10607 năm 2013)')
    for run in para.runs:
        run.font.name = 'Arial'

    para = doc.add_paragraph()
    para.add_run('Sở hữu nước ngoài: ').bold = True
    para.add_run('100% sở hữu nước ngoài được phép (RA 10881)')
    for run in para.runs:
        run.font.name = 'Arial'

    doc.add_paragraph()

    # Licensing requirements
    add_formatted_heading(doc, '3.2 Yêu Cầu Cấp Phép Môi Giới', level=2)

    capital_data = [
        ['Môi giới bảo hiểm', 'PHP 20 triệu', 'PHP 20 triệu', 'Có'],
        ['Môi giới bảo hiểm + tái bảo hiểm', 'PHP 50 triệu', 'PHP 50 triệu', 'Có']
    ]

    headers = ['Loại môi giới', 'Vốn điều lệ tối thiểu', 'Giá trị tài sản ròng', 'Cán bộ môi giới đủ tiêu chuẩn']
    add_table(doc, capital_data, headers)

    doc.add_paragraph()

    para = doc.add_paragraph()
    para.add_run('Thời hạn giấy phép: ').bold = True
    para.add_run('Gia hạn mỗi 3 năm')
    for run in para.runs:
        run.font.name = 'Arial'

    doc.add_paragraph()

    # Compulsory insurance
    add_formatted_heading(doc, '3.3 Bảo Hiểm Bắt Buộc', level=2)

    compulsory_data = [
        ['CTPL (Xe cơ giới)', 'Điều 374 Bộ luật Bảo hiểm', 'PHP 200,000 (tăng từ PHP 100,000 năm 2024)', 'Tất cả chủ xe', 'Hàng năm'],
        ['PAMI (PUV)', 'LTFRB MC 2015-028', 'PHP 200,000 tử vong, PHP 2,000-400,000 thương tích', 'Xe công cộng', 'Hàng năm'],
        ['Bảo hiểm cháy', 'Yêu cầu LGU', 'Theo chính sách', 'Doanh nghiệp (giấy phép kinh doanh)', 'Theo yêu cầu'],
        ['Hàng hải/Hàng hóa', 'Điều 92 Đạo luật Bảo hiểm', 'Theo hàng hóa', 'Nhập khẩu/Xuất khẩu (CIF 1.5%)', 'Theo lô hàng']
    ]

    headers = ['Loại bảo hiểm', 'Cơ sở pháp lý', 'Phạm vi bảo hiểm', 'Đối tượng bắt buộc', 'Thời hạn']
    add_table(doc, compulsory_data, headers)

    doc.add_paragraph()

    # EV market and EVIDA
    add_formatted_heading(doc, '3.4 Thị Trường EV và Luật EVIDA', level=2)

    para = doc.add_paragraph()
    para.add_run('Luật EVIDA (RA 11697): ').bold = True
    para.add_run('Có hiệu lực từ tháng 4/2022, cung cấp các ưu đãi toàn diện cho ngành công nghiệp xe điện')
    for run in para.runs:
        run.font.name = 'Arial'

    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run('Các ưu đãi chính:').bold = True
    for run in para.runs:
        run.font.name = 'Arial'

    incentives = [
        'Miễn "number coding" (biện pháp giảm tắc nghẽn đường bộ)',
        'Ưu tiên đăng ký và gia hạn tại LTO',
        'Giảm 30% thuế tiêu thụ đặc biệt',
        'Miễn thuế tiêu thụ đặc biệt cho xe điện nhập khẩu (đến 2028)',
        'Yêu cầu 5% đội xe của công ty phải là xe điện - tạo thị trường đảm bảo cho doanh số EV'
    ]

    for incentive in incentives:
        para = doc.add_paragraph(incentive, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.add_paragraph()

    para = doc.add_paragraph()
    para.add_run('Tăng trưởng thị trường: ').bold = True
    para.add_run('Doanh số EV tăng 80% năm 2024, hơn 20,000 xe đã đăng ký')
    for run in para.runs:
        run.font.name = 'Arial'

    doc.add_paragraph()

    # Vinfast operations
    add_formatted_heading(doc, '3.5 Hoạt Động Vinfast Hiện Tại', level=2)

    para = doc.add_paragraph()
    para.add_run('Mở showroom: ').bold = True
    para.add_run('Tháng 7/2024 - VinFast Aseana, VinFast EDSA, VinFast Alabang')
    for run in para.runs:
        run.font.name = 'Arial'

    para = doc.add_paragraph()
    para.add_run('Các mẫu xe: ').bold = True
    para.add_run('VF 5 (SUV cỡ nhỏ), VF 7 (SUV cỡ trung), VF 9 (SUV cỡ lớn)')
    for run in para.runs:
        run.font.name = 'Arial'

    para = doc.add_paragraph()
    para.add_run('Đối tác chiến lược: ').bold = True
    para.add_run('BDO Unibank (giải pháp tài chính bao gồm dịch vụ bảo hiểm), V-Green (cơ sở hạ tầng sạc), Green GSM Philippines (dịch vụ điện di chuyển)')
    for run in para.runs:
        run.font.name = 'Arial'

    para = doc.add_paragraph()
    para.add_run('Kế hoạch mở rộng: ').bold = True
    para.add_run('Hàng chục đại lý được lên kế hoạch trên khắp Philippines năm 2024')
    for run in para.runs:
        run.font.name = 'Arial'

    doc.add_paragraph()

    # Market size
    add_formatted_heading(doc, '3.6 Quy Mô Thị Trường', level=2)

    para = doc.add_paragraph()
    para.add_run('Thị trường phi nhân thọ: ').bold = True
    para.add_run('USD 2.11 tỷ (2024), dự báo đạt USD 4.27 tỷ vào 2034 (CAGR 7.3%)')
    for run in para.runs:
        run.font.name = 'Arial'

    para = doc.add_paragraph()
    para.add_run('Bảo hiểm xe cơ giới: ').bold = True
    para.add_run('41.08% phí bảo hiểm ròng (PHP 21.83 tỷ năm 2024)')
    for run in para.runs:
        run.font.name = 'Arial'

    para = doc.add_paragraph()
    para.add_run('Công ty bảo hiểm phi nhân thọ lớn nhất: ').bold = True
    para.add_run('Pioneer Insurance & Surety Corporation (PHP 15.6 tỷ năm 2023), Malayan Insurance Company, Inc.')
    for run in para.runs:
        run.font.name = 'Arial'

    doc.add_paragraph()

    # Recommendations
    add_formatted_heading(doc, '3.7 Khuyến Nghị cho Vinfast/BeeVN tại Philippines', level=2)

    recommendations = [
        'Xác minh tình trạng giấy phép môi giới hiện tại và ngày gia hạn (chu kỳ 3 năm)',
        'Đảm bảo tuân thủ vốn điều lệ PHP 20 triệu và giá trị tài sản ròng',
        'Phát triển các sản phẩm bảo hiểm chuyên biệt cho EV với các công ty bảo hiểm hàng đầu (Pioneer, Malayan, BPI/MS)',
        'Tận dụng cơ hội thị trường đội xe/TNVS thông qua chương trình hiện đại hóa PUV',
        'Hỗ trợ mở rộng mạng lưới đại lý Vinfast với các gói bảo hiểm toàn diện',
        'Định vị là môi giới bảo hiểm EV hàng đầu tại Philippines',
        'Xây dựng chuyên môn về bảo hiểm cơ sở hạ tầng sạc (trách nhiệm, tài sản, an ninh mạng)'
    ]

    for i, rec in enumerate(recommendations, 1):
        para = doc.add_paragraph(f'{i}. {rec}')
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.add_page_break()

def create_comparative_analysis(doc):
    """Create comparative analysis section"""
    add_formatted_heading(doc, 'PHẦN 4: SO SÁNH BA THỊ TRƯỜNG', level=1)

    add_formatted_heading(doc, '4.1 Ma Trận Yêu Cầu Quản Lý', level=2)

    regulatory_data = [
        ['DRC', 'ARCA', 'Không', 'SA địa phương bắt buộc', 'Chưa xác định', 'INSS nhà nước (1.5-3%)'],
        ['Kazakhstan', 'ARDFM (NBK)', 'EAEU', 'Chi nhánh/Công ty con', 'Khác nhau', 'Luật 30 (bắt buộc)'],
        ['Philippines', 'Insurance Commission', 'Không', '100% nước ngoài OK', 'PHP 20M', 'SSS/GSIS (chính phủ)']
    ]

    headers = ['Thị trường', 'Cơ quan', 'Khung khu vực', 'Hình thức pháp lý', 'Vốn tối thiểu', 'Bồi thường lao động']
    add_table(doc, regulatory_data, headers)

    doc.add_paragraph()
    add_formatted_heading(doc, '4.2 Ma Trận Sản Phẩm Bảo Hiểm cho Hoạt Động Vinfast', level=2)

    product_data = [
        ['Trách nhiệm xe cơ giới', 'BẮT BUỘC', 'BẮT BUỘC (OSAGO)', 'BẮT BUỘC (CTPL PHP 200K)'],
        ['Xe cơ giới toàn diện', 'Khuyến nghị', 'Khuyến nghị', 'Khuyến nghị'],
        ['Tài sản/Tất cả rủi ro', 'Khuyến nghị', 'Khuyến nghị', 'Khuyến nghị'],
        ['Bồi thường lao động', 'BẮT BUỘC (INSS)', 'BẮT BUỘC', 'BẮT BUỘC (SSS/GSIS)'],
        ['Trách nhiệm môi trường', 'Chưa rõ', 'BẮT BUỘC (Loại I)', 'Theo từng trường hợp'],
        ['Hàng hóa/Hàng hải', 'BẮT BUỘC (công ty DRC)', 'Khuyến nghị', 'Khuyến nghị (CIF 1.5%)'],
        ['Xây dựng (CAR)', 'BẮT BUỘC (10 năm)', 'Khuyến nghị', 'Khuyến nghị'],
        ['Trách nhiệm sản phẩm', 'Khuyến nghị', 'Khuyến nghị', 'Khuyến nghị'],
        ['Gián đoạn kinh doanh', 'Khuyến nghị', 'Khuyến nghị', 'Khuyến nghị']
    ]

    headers = ['Loại bảo hiểm', 'DRC', 'Kazakhstan', 'Philippines']
    add_table(doc, product_data, headers)

    doc.add_paragraph()
    add_formatted_heading(doc, '4.3 Ma Trận Cơ Hội Thị Trường EV', level=2)

    ev_data = [
        ['DRC', 'Tối thiểu', 'Không cụ thể', 'Chưa phát triển', 'Nhà máy pin'],
        ['Kazakhstan', '19,000+ (H1 2025)', 'Đến 31/12/2025', 'Chưa phát triển', 'Mở rộng dự kiến'],
        ['Philippines', '20,000+ (2024)', 'EVIDA (5% đội xe)', 'Mới nổi', 'Đại lý hoạt động từ 7/2024']
    ]

    headers = ['Thị trường', 'Quy mô thị trường EV', 'Ưu đãi chính phủ', 'Sản phẩm bảo hiểm EV', 'Tình trạng hoạt động Vinfast']
    add_table(doc, ev_data, headers)

    doc.add_paragraph()
    add_formatted_heading(doc, '4.4 Phân Tích Điểm Mạnh-Yếu-Cơ Hội-Thách Thức', level=2)

    # DRC SWOT
    para = doc.add_paragraph()
    para.add_run('DRC:').bold = True
    for run in para.runs:
        run.font.name = 'Arial'

    swot_drc = [
        ('Điểm mạnh', 'Tăng trưởng thị trường nhanh (5 lần), tự do hóa gần đây'),
        ('Điểm yếu', 'Rủi ro chính trị cao, yêu cầu bảo hiểm hàng nhập khẩu địa phương bắt buộc'),
        ('Cơ hội', 'Nhà máy pin, thị trường chưa phát triển'),
        ('Thách thức', 'Yêu cầu cấp phép không rõ ràng, thiếu dữ liệu')
    ]

    for category, content in swot_drc:
        para = doc.add_paragraph()
        run = para.add_run(f'• {category}: ')
        run.bold = True
        run.font.name = 'Arial'
        run = para.add_run(content)
        run.font.name = 'Arial'

    doc.add_paragraph()

    # Kazakhstan SWOT
    para = doc.add_paragraph()
    para.add_run('Kazakhstan:').bold = True
    for run in para.runs:
        run.font.name = 'Arial'

    swot_kz = [
        ('Điểm mạnh', 'Môi trường quản lý mạnh, tăng trưởng thị trường EV nhanh, ưu đãi chính phủ'),
        ('Điểm yếu', 'Cấp phép EAEU không tự động công nhận'),
        ('Cơ hội', 'Thị trường bảo hiểm EV chưa phát triển - cơ hội chuyên môn hóa lớn'),
        ('Thách thức', 'Cần nghiên cứu sơ cấp về các sản phẩm cụ thể')
    ]

    for category, content in swot_kz:
        para = doc.add_paragraph()
        run = para.add_run(f'• {category}: ')
        run.bold = True
        run.font.name = 'Arial'
        run = para.add_run(content)
        run.font.name = 'Arial'

    doc.add_paragraph()

    # Philippines SWOT
    para = doc.add_paragraph()
    para.add_run('Philippines:').bold = True
    for run in para.runs:
        run.font.name = 'Arial'

    swot_ph = [
        ('Điểm mạnh', '100% sở hữu nước ngoài, hoạt động Vinfast hiện có, luật EVIDA mạnh'),
        ('Điểm yếu', 'Thị trường cạnh tranh, yêu cầu gia hạn 3 năm'),
        ('Cơ hội', 'Hiện đại hóa PUV, yêu cầu 5% đội xe, mở rộng đại lý'),
        ('Thách thức', 'Bảo hiểm EV đắt hơn, cần giáo dục khách hàng')
    ]

    for category, content in swot_ph:
        para = doc.add_paragraph()
        run = para.add_run(f'• {category}: ')
        run.bold = True
        run.font.name = 'Arial'
        run = para.add_run(content)
        run.font.name = 'Arial'

    doc.add_page_break()

def create_action_plan(doc):
    """Create action plan section"""
    add_formatted_heading(doc, 'PHẦN 5: KẾ HOẠCH HÀNH ĐỘNG', level=1)

    add_formatted_heading(doc, 'Giai Đoạn 1: Ngay Lập Tức (0-30 ngày)', level=2)

    phase1_actions = [
        ('Philippines', 'Xác minh tình trạng giấy phép môi giới; đảm bảo tuân thủ hỗ trợ Vinfast'),
        ('Tất cả thị trường', 'Thuê tư vấn pháp lý địa phương cho từng quyền tài phán'),
        ('Tất cả thị trường', 'Liên hệ trực tiếp với cơ quan quản lý để làm rõ yêu cầu cấp phép')
    ]

    for market, action in phase1_actions:
        para = doc.add_paragraph()
        run = para.add_run(f'{market}: ')
        run.bold = True
        run.font.name = 'Arial'
        run = para.add_run(action)
        run.font.name = 'Arial'

    doc.add_paragraph()
    add_formatted_heading(doc, 'Giai Đoạn 2: Ngắn Hạn (30-90 ngày)', level=2)

    phase2_actions = [
        ('DRC', 'Bắt đầu tư vấn ARCA; khởi động quy trình thành lập công ty'),
        ('Kazakhstan', 'Bắt đầu đơn xin chi nhánh qua elicense.kz'),
        ('Philippines', 'Phát triển các sản phẩm bảo hiểm chuyên biệt EV với các công ty bảo hiểm')
    ]

    for market, action in phase2_actions:
        para = doc.add_paragraph()
        run = para.add_run(f'{market}: ')
        run.bold = True
        run.font.name = 'Arial'
        run = para.add_run(action)
        run.font.name = 'Arial'

    doc.add_paragraph()
    add_formatted_heading(doc, 'Giai Đoạn 3: Trung Hạn (90-180 ngày)', level=2)

    phase3_actions = [
        ('DRC', 'Hoàn thành cấp phép môi giới; thiết lập văn phòng tại Kinshasa'),
        ('Kazakhstan', 'Có được giấy phép hoạt động; thuê nhân viên địa phương'),
        ('Tất cả thị trường', 'Ra mắt chương trình bảo hiểm Vinfast phối hợp')
    ]

    for market, action in phase3_actions:
        para = doc.add_paragraph()
        run = para.add_run(f'{market}: ')
        run.bold = True
        run.font.name = 'Arial'
        run = para.add_run(action)
        run.font.name = 'Arial'

    doc.add_paragraph()
    add_formatted_heading(doc, 'Giai Đoạn 4: Dài Hạn (180+ ngày)', level=2)

    phase4_actions = [
        ('Tất cả thị trường', 'Định vị là chuyên gia bảo hiểm EV'),
        ('Khu vực', 'Phát triển khả năng quản lý rủi ro xuyên biên giới'),
        ('Đổi mới', 'Tạo ra các sản phẩm dựa trên viễn thông cho đội xe EV')
    ]

    for category, action in phase4_actions:
        para = doc.add_paragraph()
        run = para.add_run(f'{category}: ')
        run.bold = True
        run.font.name = 'Arial'
        run = para.add_run(action)
        run.font.name = 'Arial'

    doc.add_page_break()

def create_data_gaps(doc):
    """Create data gaps section"""
    add_formatted_heading(doc, 'PHẦN 6: KHOẢNG TRỐNG DỮ LIỆU VÀ NGHIÊN CỨU TIẾP THEO', level=1)

    add_formatted_heading(doc, '6.1 DRC - Cần Nghiên Cứu Sơ Cấp', level=2)

    drc_gaps = [
        'Yêu cầu bảo hiểm trách nhiệm môi trường (xác nhận với ARCA)',
        'Yêu cầu vốn cụ thể cho môi giới',
        'Yêu cầu đối tác địa phương cho môi giới nước ngoài',
        'Thuế suất phí bảo hiểm',
        'Yêu cầu ngôn ngữ chính sách (Pháp/Anh)',
        'Thời gian giải quyết khiếu nại tiêu chuẩn thị trường',
        'Công suất thị trường cho rủi ro công nghiệp lớn'
    ]

    for gap in drc_gaps:
        para = doc.add_paragraph(gap, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_formatted_heading(doc, '6.2 Kazakhstan - Cần Nghiên Cứu Sơ Cấp', level=2)

    kz_gaps = [
        'Tính khả dụng của các sản phẩm bảo hiểm cụ thể cho EV',
        'Yêu cầu bắt buộc bảo hiểm CAR',
        'Quy định cụ thể cho đội xe/taxi',
        'Cấu trúc hoa hồng môi giới',
        'Thời gian cấp phép ARDFM',
        'Tiêu chí bảo lãnh cho rủi ro EV',
        'Sắp xếp tái bảo hiểm cho hoạt động EV'
    ]

    for gap in kz_gaps:
        para = doc.add_paragraph(gap, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_formatted_heading(doc, '6.3 Philippines - Cần Nghiên Cứu Sơ Cấp', level=2)

    ph_gaps = [
        'Yêu cầu CPD cho môi giới (nếu có)',
        'Yêu cầu bảo hiểm bắt buộc cụ thể cho đại lý',
        'Chi tiết sản phẩm bảo hiểm EV từ các công ty bảo hiểm cụ thể',
        'Yêu cầu bảo hiểm trạm sạc',
        'Giá và điều khoản bảo vệ so sánh',
        'Thực hành ngành công nghiệp ô tô (Toyota, Honda, Nissan)',
        'Chiến lược bảo hiểm nhãn hiệu EV khác (BYD, Tesla, MG)'
    ]

    for gap in ph_gaps:
        para = doc.add_paragraph(gap, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_formatted_heading(doc, '6.4 Các Bên Liên Quan Chính để Tư Vấn', level=2)

    stakeholders_data = [
        ['DRC', 'ARCA, Tư vấn pháp lý Congo, SFA, Rawsur, SONAS', 'arca.cd'],
        ['Kazakhstan', 'ARDFM, IC Centras, Eurasia Insurance, Halyk', 'www.elicense.kz'],
        ['Philippines', 'Insurance Commission, Pioneer, Malayan, BPI/MS, BDO', 'www.insurance.gov.ph']
    ]

    headers = ['Thị trường', 'Các bên liên quan', 'Website chính']
    add_table(doc, stakeholders_data, headers)

    doc.add_page_break()

def create_appendices(doc):
    """Create appendices section"""
    add_formatted_heading(doc, 'PHỤ LỤC A: LIÊN HỆ VÀ NGUỒN TÀI LIỆU', level=1)

    add_formatted_heading(doc, 'DRC - Cộng Hòa Dân Chủ Congo', level=2)

    drc_contacts = [
        ('ARCA', 'arca.cd', 'Cơ quan quản lý'),
        ('SONAS', 'www.sonas.cd, +243 81 970 0904', 'Công ty bảo hiểm nhà nước'),
        ('SFA', 'Không có trang web công khai', 'Công ty dẫn đầu thị trường')
    ]

    for org, contact, note in drc_contacts:
        para = doc.add_paragraph()
        run = para.add_run(f'{org}: ')
        run.bold = True
        run.font.name = 'Arial'
        run = para.add_run(f'{contact} ({note})')
        run.font.name = 'Arial'

    doc.add_paragraph()
    add_formatted_heading(doc, 'Kazakhstan', level=2)

    kz_contacts = [
        ('ARDFM', 'gov.kz/memleket/entities/ardfm', 'Cơ quan quản lý chính'),
        ('E-licensing', 'www.elicense.kz', 'Cổng cấp phép điện tử'),
        ('IC Centras', 'Không rõ', 'Công ty dẫn đầu thị trường'),
        ('Eurasia Insurance', 'Không rõ', 'Công ty phi nhân thọ lớn')
    ]

    for org, contact, note in kz_contacts:
        para = doc.add_paragraph()
        run = para.add_run(f'{org}: ')
        run.bold = True
        run.font.name = 'Arial'
        run = para.add_run(f'{contact} ({note})')
        run.font.name = 'Arial'

    doc.add_paragraph()
    add_formatted_heading(doc, 'Philippines', level=2)

    ph_contacts = [
        ('Insurance Commission', 'www.insurance.gov.ph', 'Cơ quan quản lý chính'),
        ('Pioneer Insurance', 'Không rõ', 'Công ty dẫn đầu thị trường'),
        ('Malayan Insurance', 'Không rõ', 'Công ty lớn thứ hai'),
        ('BDO Unibank', 'www.bdo.com.ph', 'Đối tác chiến lược Vinfast')
    ]

    for org, contact, note in ph_contacts:
        para = doc.add_paragraph()
        run = para.add_run(f'{org}: ')
        run.bold = True
        run.font.name = 'Arial'
        run = para.add_run(f'{contact} ({note})')
        run.font.name = 'Arial'

    doc.add_paragraph()
    add_formatted_heading(doc, 'PHỤ LỤC B: NGUỒN NGHIÊN CỨU', level=1)

    para = doc.add_paragraph()
    para.add_run('Báo cáo này được biên soạn từ các nguồn thứ cấp đáng tin cậy bao gồm:').font.name = 'Arial'

    sources = [
        'Các ấn phẩm chính phủ chính thức (ARCA, ARDFM, Insurance Commission)',
        'Cơ quan thương mại quốc tế (US Trade.gov)',
        'Các ấn phẩm ngành bảo hiểm (Atlas Magazine, Insurance Journal, Reinsurance News)',
        'Nền tảng phân tích pháp lý (HG.org, Lexology, Chambers)',
        'Các công ty nghiên cứu thị trường (GlobalData, Axco, Statista)',
        'Cơ quan an sinh xã hội (SSA.gov)',
        'Báo cáo chính thức của Vinfast'
    ]

    for source in sources:
        para = doc.add_paragraph(source, style='List Bullet')
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run('Tất cả dữ liệu trong vòng 12 tháng (2025-2026). Các báo cáo nghiên cứu chi tiết đầy đủ có sẵn bằng tiếng Anh.').font.name = 'Arial'

def create_document():
    """Main function to create the complete document"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Create all sections
    create_cover_page(doc)
    create_executive_summary(doc)
    create_drc_section(doc)
    create_kazakhstan_section(doc)
    create_philippines_section(doc)
    create_comparative_analysis(doc)
    create_action_plan(doc)
    create_data_gaps(doc)
    create_appendices(doc)

    # Add page numbers
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = "Trang "

    # Save document
    output_path = '/Users/nntrvy/osr-company/os-research-engine/research-outputs/26011-vinfast-insurance-markets/04.02.26-beevn-bao-cao-thi-truong-bao-hiem.docx'
    doc.save(output_path)
    print(f"Document created successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    create_document()
