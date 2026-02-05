#!/usr/bin/env python3
"""
Create .docx files with OS Research standard formatting.

Based on: Thử việc - Content Strategist - Đào Thanh Tâm - 02_02_2026.docx

Font: IBM Plex Sans
Page: A4
Margins: 22mm top/left, 20mm right/bottom
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class OSRDocument:
    """Creates documents with OS Research standard formatting."""

    FONT_NAME = "IBM Plex Sans"
    FONT_SIZE = Pt(12)
    TITLE_SIZE = Pt(26)
    HEADING_SIZE = Pt(14)

    # A4 page size
    PAGE_WIDTH = Inches(8.27)
    PAGE_HEIGHT = Inches(11.69)

    # Margins
    MARGIN_TOP = Cm(2.2)
    MARGIN_BOTTOM = Cm(2.0)
    MARGIN_LEFT = Cm(2.2)
    MARGIN_RIGHT = Cm(2.0)

    # Spacing
    SPACE_BEFORE = Pt(12)
    SPACE_AFTER = Pt(12)
    SPACE_AFTER_SMALL = Pt(6)
    LINE_SPACING = 1.0

    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_default_style()

    def _setup_page(self):
        """Configure page size and margins."""
        section = self.doc.sections[0]
        section.page_width = self.PAGE_WIDTH
        section.page_height = self.PAGE_HEIGHT
        section.top_margin = self.MARGIN_TOP
        section.bottom_margin = self.MARGIN_BOTTOM
        section.left_margin = self.MARGIN_LEFT
        section.right_margin = self.MARGIN_RIGHT

    def _setup_default_style(self):
        """Configure default paragraph style."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.FONT_NAME
        font.size = self.FONT_SIZE

        # Set East Asian font (for Vietnamese diacritics)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_NAME)

    def _set_font(self, run, font_name=None, size=None, bold=False, italic=False, underline=False):
        """Apply font formatting to a run."""
        font_name = font_name or self.FONT_NAME
        size = size or self.FONT_SIZE

        run.font.name = font_name
        run.font.size = size
        run.bold = bold
        run.italic = italic
        run.underline = underline

        # Set East Asian font
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)

    def _set_paragraph_format(self, para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               space_before=None, space_after=None):
        """Apply paragraph formatting."""
        space_before = space_before if space_before is not None else self.SPACE_BEFORE
        space_after = space_after if space_after is not None else self.SPACE_AFTER

        pf = para.paragraph_format
        pf.alignment = alignment
        pf.space_before = space_before
        pf.space_after = space_after
        pf.line_spacing = self.LINE_SPACING

    def add_header(self):
        """Add Vietnam official document header."""
        # National title
        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, WD_ALIGN_PARAGRAPH.CENTER)
        run = para.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
        self._set_font(run)

        # Motto
        para.add_run("\n")
        run = para.add_run("Độc lập – Tự do – Hạnh phúc")
        self._set_font(run, italic=True)

        # Divider
        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, WD_ALIGN_PARAGRAPH.CENTER)
        run = para.add_run("---------o0o---------")
        self._set_font(run)

    def add_date_line(self, location="TPHCM"):
        """Add date line (right-aligned)."""
        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, WD_ALIGN_PARAGRAPH.RIGHT)
        run = para.add_run(f"{location}, ngày …….. tháng …….. năm 20……..")
        self._set_font(run, italic=True)

    def add_title(self, text, doc_number=None):
        """Add document title (centered, bold, large)."""
        # Title
        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, WD_ALIGN_PARAGRAPH.CENTER)
        run = para.add_run(text)
        self._set_font(run, size=self.TITLE_SIZE, bold=True)

        # Document number
        if doc_number:
            para = self.doc.add_paragraph()
            self._set_paragraph_format(para, WD_ALIGN_PARAGRAPH.CENTER)
            run = para.add_run("Số: ")
            self._set_font(run, italic=True)
            run = para.add_run(doc_number)
            self._set_font(run, bold=True)

    def add_section_heading(self, text, number=None):
        """Add a section heading (bold)."""
        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT)

        if number:
            run = para.add_run(f"Điều {number}: ")
            self._set_font(run, bold=True)

        run = para.add_run(text)
        self._set_font(run, bold=True)

    def add_paragraph(self, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      bold=False, italic=False, space_after=None):
        """Add a regular paragraph."""
        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, alignment, space_after=space_after)
        run = para.add_run(text)
        self._set_font(run, bold=bold, italic=italic)
        return para

    def add_bullet(self, text, indent_level=0):
        """Add a bullet point."""
        prefix = "    " * indent_level + "- "
        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   space_after=self.SPACE_AFTER_SMALL)
        run = para.add_run(prefix + text)
        self._set_font(run)
        return para

    def add_numbered_item(self, number, text):
        """Add a numbered item."""
        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   space_after=self.SPACE_AFTER_SMALL)
        run = para.add_run(f"{number}. ")
        self._set_font(run, bold=True)
        run = para.add_run(text)
        self._set_font(run)
        return para

    def add_signature_block(self, left_title, left_name, right_title, right_name):
        """Add two-column signature block."""
        # Create table for alignment
        table = self.doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(3)
            row.cells[1].width = Inches(3)

        # Left side
        cell = table.cell(0, 0)
        para = cell.paragraphs[0]
        self._set_paragraph_format(para, WD_ALIGN_PARAGRAPH.CENTER)
        run = para.add_run(left_title)
        self._set_font(run, bold=True)

        cell = table.cell(1, 0)
        para = cell.paragraphs[0]
        self._set_paragraph_format(para, WD_ALIGN_PARAGRAPH.CENTER)
        run = para.add_run("(Ký và ghi rõ họ tên)")
        self._set_font(run, italic=True)

        cell = table.cell(3, 0)
        para = cell.paragraphs[0]
        self._set_paragraph_format(para, WD_ALIGN_PARAGRAPH.CENTER)
        run = para.add_run(left_name)
        self._set_font(run, bold=True)

        # Right side
        cell = table.cell(0, 1)
        para = cell.paragraphs[0]
        self._set_paragraph_format(para, WD_ALIGN_PARAGRAPH.CENTER)
        run = para.add_run(right_title)
        self._set_font(run, bold=True)

        cell = table.cell(1, 1)
        para = cell.paragraphs[0]
        self._set_paragraph_format(para, WD_ALIGN_PARAGRAPH.CENTER)
        run = para.add_run("(Ký và ghi rõ họ tên)")
        self._set_font(run, italic=True)

        cell = table.cell(3, 1)
        para = cell.paragraphs[0]
        self._set_paragraph_format(para, WD_ALIGN_PARAGRAPH.CENTER)
        run = para.add_run(right_name)
        self._set_font(run, bold=True)

    def add_empty_lines(self, count=1):
        """Add empty lines."""
        for _ in range(count):
            para = self.doc.add_paragraph()
            self._set_paragraph_format(para, space_before=Pt(0), space_after=Pt(0))

    def save(self, filename):
        """Save the document."""
        self.doc.save(filename)
        print(f"Document saved: {filename}")


def create_sample_contract():
    """Create a sample probation contract."""
    doc = OSRDocument()

    # Header
    doc.add_header()
    doc.add_date_line("TPHCM")

    # Title
    doc.add_title("HỢP ĐỒNG THỬ VIỆC", "01/2026/HDTV-OSR")

    # Legal basis
    doc.add_paragraph("Căn cứ:", bold=True)
    doc.add_bullet("Bộ luật Lao động số 45/2019/QH14 ngày 20/11/2019;")
    doc.add_bullet("Nghị định 145/2020/NĐ-CP ngày 14/12/2020;")

    # Parties
    doc.add_paragraph("Hôm nay, chúng tôi gồm:")
    doc.add_section_heading("BÊN A – NGƯỜI SỬ DỤNG LAO ĐỘNG", number=None)
    doc.add_paragraph("Tên doanh nghiệp: CÔNG TY TNHH OS RESEARCH")
    doc.add_paragraph("Địa chỉ: [Địa chỉ công ty]")
    doc.add_paragraph("Đại diện: Ông/Bà [Tên]    Chức vụ: Giám đốc")

    doc.add_empty_lines(1)

    doc.add_section_heading("BÊN B – NGƯỜI LAO ĐỘNG", number=None)
    doc.add_paragraph("Họ và tên: [Họ tên nhân viên]")
    doc.add_paragraph("Ngày sinh: [DD/MM/YYYY]    Giới tính: [Nam/Nữ]")
    doc.add_paragraph("CCCD số: [Số CCCD]")
    doc.add_paragraph("Địa chỉ thường trú: [Địa chỉ]")

    doc.add_empty_lines(1)

    # Agreement
    doc.add_paragraph("Hai bên cùng thỏa thuận ký kết Hợp đồng thử việc với các điều khoản sau:")

    # Articles
    doc.add_section_heading("Công việc và địa điểm làm việc", number=1)
    doc.add_paragraph("Chức danh: [Vị trí]")
    doc.add_paragraph("Địa điểm làm việc: [Địa điểm]")

    doc.add_section_heading("Thời hạn thử việc", number=2)
    doc.add_paragraph("Thời gian thử việc: [X] ngày, từ ngày [DD/MM/YYYY] đến ngày [DD/MM/YYYY]")

    doc.add_section_heading("Lương thử việc", number=3)
    doc.add_paragraph("Mức lương thử việc: [X.XXX.XXX] VND/tháng (bằng 85% lương chính thức)")

    doc.add_empty_lines(2)

    # Signatures
    doc.add_signature_block(
        "ĐẠI DIỆN BÊN A", "[Tên người đại diện]",
        "BÊN B", "[Tên người lao động]"
    )

    doc.save("sample_contract.docx")


if __name__ == "__main__":
    create_sample_contract()
    print("\nUsage example:")
    print("  from create_osr_docx import OSRDocument")
    print("  doc = OSRDocument()")
    print("  doc.add_header()")
    print("  doc.add_title('TÊN VĂN BẢN')")
    print("  doc.add_paragraph('Nội dung...')")
    print("  doc.save('output.docx')")
