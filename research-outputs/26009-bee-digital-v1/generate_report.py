#!/usr/bin/env python3
"""
Generate comprehensive market research report for Electrik/BeeVN
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add a heading with custom formatting"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    doc.add_paragraph()
    return table

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('ELECTRIK MARKET RESEARCH REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Comprehensive Analysis for BeeVN Insurance Technology Platform')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph(f'Prepared by: OS Research Company\nDate: {datetime.now().strftime("%B %d, %Y")}\nVersion: 1.0')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    add_heading(doc, 'TABLE OF CONTENTS', 1)
    toc_items = [
        '1. Executive Summary',
        '2. Market Overview',
        '3. Current State Analysis',
        '4. Stakeholder Analysis',
        '5. Competitive Landscape',
        '6. Customer Interview Insights',
        '7. Validated vs Invalidated Hypotheses',
        '8. Opportunity Sizing',
        '9. Strategic Recommendations',
        '10. Implementation Roadmap',
        '11. Appendices'
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, '1. EXECUTIVE SUMMARY', 1)

    add_heading(doc, '1.1 Project Background', 2)
    doc.add_paragraph(
        'The Electrik project began in July 2025 as an embedded insurance distribution concept '
        'for VinFast electric vehicles. Through rigorous validation over three months, including '
        'customer interviews with insurers, OEMs, and garages, the project pivoted to become a '
        'claims processing platform for the Collision Market.'
    )

    add_heading(doc, '1.2 Key Finding', 2)
    doc.add_paragraph(
        'Target Market: Collision Market (garages + insurers), NOT car owners. '
        'Vietnamese car owners are already well-served compared to global standards. '
        'The real pain points exist with garages (cash flow, paperwork) and insurers '
        '(digitization pressure, operational efficiency).'
    )

    add_heading(doc, '1.3 Market Opportunity', 2)
    doc.add_paragraph('• Motor vehicle claims market: ~1,000 billion VND/year')
    doc.add_paragraph('• VinFast EV sales 2025: 175,099 vehicles (record)')
    doc.add_paragraph('• Insurance distribution potential: ~29 billion VND/year at 15% penetration')
    doc.add_paragraph('• Fraud estimated at 15% of claims costs industry-wide')

    add_heading(doc, '1.4 Recommended Entry Points', 2)
    doc.add_paragraph('1. AI Call Center – Explicit demand from insurer interviews')
    doc.add_paragraph('2. Garage Portal – High pain, low politics, clear ROI')
    doc.add_paragraph('3. Claims Triage System – Meets insurer 40% AI processing KPI')

    doc.add_page_break()

    # 2. Market Overview
    add_heading(doc, '2. MARKET OVERVIEW', 1)

    add_heading(doc, '2.1 Vietnam Auto Insurance Market', 2)

    add_table(doc,
        ['Metric', '2024 Value', 'Notes'],
        [
            ['Total cars sold', '~494,000 vehicles', 'VAMA data'],
            ['EVs sold', '87,000+ vehicles', '~18% market share'],
            ['VinFast EV share', '90%+', 'Dominant player'],
            ['VinFast 2025 sales', '175,099 EVs', 'Record year, nearly doubled 2024'],
        ]
    )

    add_heading(doc, '2.2 Market Projections', 2)

    add_table(doc,
        ['Year', 'Projection', 'Source'],
        [
            ['2025', '$2.93 billion (EV market)', 'Mordor Intelligence'],
            ['2028', '1 million EVs on road', 'VAMA'],
            ['2030', '$6.69 billion (EV market)', 'Mordor Intelligence'],
            ['2040', '3.5 million EVs on road', 'VAMA'],
        ]
    )

    doc.add_paragraph('Compound Annual Growth Rate (CAGR): ~18%')

    add_heading(doc, '2.3 Claims Market Size', 2)
    doc.add_paragraph(
        'The motor vehicle claims market in Vietnam represents approximately 1,000 billion VND '
        'in annual payouts. Major insurers include PTI, MIC, PJICO, BSH, and DBV. '
        'Fraud is estimated to account for approximately 15% of total claims costs, '
        'representing significant industry-wide losses.'
    )

    add_heading(doc, '2.4 VinFast 2025 Performance', 2)

    add_table(doc,
        ['Model', 'Units Sold', 'Notes'],
        [
            ['VF 3', '44,585', 'Best-seller'],
            ['VF 5', '43,913', '#2 best-seller'],
            ['Limo Green', '27,127', 'B2B/taxi segment'],
            ['VF 6', '23,291', 'Mid-range'],
            ['Herio Green', '12,568', 'B2B segment'],
            ['VF 7', '9,653', 'Premium'],
        ]
    )

    doc.add_paragraph(
        'VinFast achieved #1 automotive brand status in Vietnam for 15 consecutive months, '
        'with December 2025 setting a record of 27,649 units delivered in a single month.'
    )

    doc.add_page_break()

    # 3. Current State Analysis
    add_heading(doc, '3. CURRENT STATE ANALYSIS', 1)

    add_heading(doc, '3.1 BeeVN Current Situation', 2)
    doc.add_paragraph(
        'BeeVN is Vingroup\'s in-house insurance broker, responsible for arranging commercial '
        'insurance for all Vingroup subsidiaries and negotiating claims payouts with insurers.'
    )

    doc.add_paragraph('Current Partnerships:')
    doc.add_paragraph('• PVI Digital: Trip insurance for GSM passengers')
    doc.add_paragraph('• Malayan: Driver savings insurance (GSM Philippines)')

    add_heading(doc, '3.2 GSM Trip Insurance Case Study', 2)
    doc.add_paragraph(
        'BeeVN partnered with PVI Digital to sell trip insurance to GSM ride-hailing passengers '
        'at 2,000 VND per trip. Initial auto opt-in drove 15 billion VND/month in revenue. '
        'After CQLTT (market regulators) required explicit opt-in consent, revenue dropped '
        'to approximately 8 billion VND/month – a 47% decline.'
    )

    add_heading(doc, '3.3 Core Issues Identified', 2)

    add_table(doc,
        ['Issue', 'Description', 'Impact'],
        [
            ['No system control', 'BeeVN depends on partners for technology', 'Cannot iterate quickly'],
            ['Slow integration', 'GSM-PVI integration requires workarounds', 'Delayed product launches'],
            ['No demand control', 'Cannot influence customer conversion', 'Revenue vulnerability'],
        ]
    )

    add_heading(doc, '3.4 Industry-Wide Pain Points', 2)
    doc.add_paragraph('• Slow claims process: 15-30 days typical resolution time')
    doc.add_paragraph('• Manual paperwork: Error-prone, human-dependent processes')
    doc.add_paragraph('• Fragmented systems: No unified customer-garage-insurer connection')
    doc.add_paragraph('• No EV-specific solution: Specialized needs not addressed')
    doc.add_paragraph('• Declining trust: Social confidence in insurance eroding')

    doc.add_page_break()

    # 4. Stakeholder Analysis
    add_heading(doc, '4. STAKEHOLDER ANALYSIS', 1)

    add_heading(doc, '4.1 Car Owners', 2)
    doc.add_paragraph('Evidence Strength: Weak to Medium')
    doc.add_paragraph()
    doc.add_paragraph(
        'Key Finding: Insurance claims are NOT a major pain point for Vietnamese car owners. '
        'Compared to foreign markets, Vietnamese customers are relatively well-served. '
        'Any solution targeting this segment would be "vitamins" rather than "painkillers."'
    )
    doc.add_paragraph()
    doc.add_paragraph('Implication: Do not lead product development with consumer-facing features.')

    add_heading(doc, '4.2 Insurers', 2)
    doc.add_paragraph('Evidence Strength: Strong')
    doc.add_paragraph()
    doc.add_paragraph('Key Findings:')
    doc.add_paragraph('• Pain is real and significant but NOT urgent – they are managing with current solutions')
    doc.add_paragraph('• Specific KPI revealed: "40% of claims must be processed fully by AI, without adjuster intervention"')
    doc.add_paragraph('• Fraud detection is NOT a priority – considered an "unwritten rule" (luật ngầm) in the industry')
    doc.add_paragraph('• AI Call Center = explicit urgent demand (partner asked: "Can you build this for me?")')
    doc.add_paragraph()
    doc.add_paragraph(
        'Profile - MIC Head of Claims: Mid-level manager with digital transformation mandate, '
        'can order IT solutions, operates software after handover. Has real pain but not desperate.'
    )

    add_heading(doc, '4.3 Garages', 2)
    doc.add_paragraph('Evidence Strength: Strong')
    doc.add_paragraph()
    doc.add_paragraph('Key Findings:')
    doc.add_paragraph('• Most important and powerful party in the claims chain')
    doc.add_paragraph('• Simple to work with, less political than other stakeholders')
    doc.add_paragraph('• Average monthly revenue: 3.6 billion VND')
    doc.add_paragraph('• Insurance customers: ~15% of revenue (~600 million VND/month)')
    doc.add_paragraph('• Insurance work is more profitable than warranty work')
    doc.add_paragraph()
    doc.add_paragraph('Pain Points:')
    doc.add_paragraph('• Document preparation is time-consuming and labor-intensive')
    doc.add_paragraph('• Insurers don\'t notify on incomplete documents – garage must chase')
    doc.add_paragraph('• DMS does not integrate with MISA accounting software – manual copying required')
    doc.add_paragraph('• Insurers don\'t accept garage photos – adjuster must visit in person')

    add_heading(doc, '4.4 VinFast/OEM', 2)
    doc.add_paragraph('Evidence Strength: Weak')
    doc.add_paragraph()
    doc.add_paragraph('Key Findings:')
    doc.add_paragraph('• Will consider integration IF insurers have a proven, reliable solution')
    doc.add_paragraph('• Will NOT co-develop from scratch when nothing is proven')
    doc.add_paragraph('• Better insurance experience is a small priority in VinFast strategy')
    doc.add_paragraph('• Only care about GSM as main customer')
    doc.add_paragraph('• Everything requires Vingroup chairman approval')
    doc.add_paragraph()
    doc.add_paragraph('Implication: Do not waste time trying to partner with VinFast until insurers are onboard.')

    doc.add_page_break()

    # 5. Competitive Landscape
    add_heading(doc, '5. COMPETITIVE LANDSCAPE', 1)

    add_heading(doc, '5.1 Local Players', 2)

    add_table(doc,
        ['Player', 'Strength', 'Gap', 'Rating'],
        [
            ['Fermion/Merimen', 'Large system, insurer relationships', 'Not EV-optimized, old-timer', '3/5'],
            ['FPT.AI', 'Strong AI (OCR/NLP)', 'Mostly chatbots, not end-to-end claims', '3/5'],
            ['AiCycle.ai', 'AI for fraud/claims, local partnerships', 'Pivoting to distribution – leaving tech gap', '3/5'],
            ['Local apps (Bảo Việt, PVI, PTI)', 'Market presence', 'Just digitized paperwork, no CX/AI', '2/5'],
            ['Global Care', 'Platform & API, acquired by Huize', 'Limited claims focus', '4/5'],
        ]
    )

    add_heading(doc, '5.2 International References', 2)

    add_table(doc,
        ['Company', 'Country', 'Strength', 'Learn For Electrik'],
        [
            ['Lemonade', 'US', 'Best-in-class CX', '"Wow" experience, trust through transparency'],
            ['FRISS', 'Netherlands', 'Fraud scoring AI', 'Flip to "Honest Score" (positive framing)'],
            ['Ravin.ai', 'Israel', 'AI damage detection', 'One-touch video claim processing'],
            ['CCC', 'US', 'Collision market platform', 'Reference model for garage-insurer connection'],
            ['Root Insurance', 'US', 'Usage-based insurance', 'Light version: km, charging, basic driving'],
            ['PasarPolis', 'Indonesia', 'Ecosystem integration', 'Plug into existing apps'],
        ]
    )

    add_heading(doc, '5.3 Blue Ocean Positioning', 2)
    doc.add_paragraph(
        'Electrik targets the upper-right quadrant of the competitive landscape: '
        'High Automation + High Customer Experience, specialized for EV ecosystem.'
    )
    doc.add_paragraph()
    doc.add_paragraph('Differentiation:')
    doc.add_paragraph('• vs. Fermion/Merimen: Higher customer experience, EV-specialized')
    doc.add_paragraph('• vs. FPT.AI: End-to-end solution, not just AI components')
    doc.add_paragraph('• vs. Insurtech apps: Deeper automation, not just pretty UI')
    doc.add_paragraph('• vs. Traditional adjusters: Technology-first, scalable')

    add_heading(doc, '5.4 Key Differentiator: "Honest Score"', 2)
    doc.add_paragraph(
        'Instead of "Fraud Score" (negative, suspicious framing), Electrik proposes "Honest Score" – '
        'a positive framing that rewards transparent customers with faster processing. '
        'Same AI technology, better customer psychology. Builds trust instead of suspicion.'
    )

    doc.add_page_break()

    # 6. Customer Interview Insights
    add_heading(doc, '6. CUSTOMER INTERVIEW INSIGHTS', 1)

    add_heading(doc, '6.1 Interviews Conducted', 2)
    doc.add_paragraph('Date: September 24-25, 2025')
    doc.add_paragraph()

    add_table(doc,
        ['Person', 'Role', 'Organization'],
        [
            ['Anh Phong', 'Head of Claims', 'MIC (insurer)'],
            ['Anh Bảo', 'Product Manager', 'XanhSM/VinFast'],
            ['Workshop Manager', 'Operations', 'VinFast GMS Gia Lâm'],
            ['Claims Team', 'Claims Processing', 'PJICO (insurer)'],
        ]
    )

    add_heading(doc, '6.2 Critical Findings', 2)

    doc.add_paragraph('Finding 1: Car Owners = NOT a Big Pain Point', style='Intense Quote')
    doc.add_paragraph(
        '"Vietnamese car owners are relatively well-served compared to foreign markets." '
        'Any solution is vitamins, not painkillers.'
    )

    doc.add_paragraph('Finding 2: Fraud Detection = NOT a Priority', style='Intense Quote')
    doc.add_paragraph(
        '"Fraud at some level is an unwritten rule (luật ngầm) in the industry." '
        'Leadership doesn\'t care because it\'s hard to quantify savings. Do NOT lead with fraud detection.'
    )

    doc.add_paragraph('Finding 3: AI Call Center = Urgent Need', style='Intense Quote')
    doc.add_paragraph(
        'During the interview, a partner immediately asked: "Can you build this solution for me?" '
        'This indicates extreme urgency and represents the best entry point for quick wins.'
    )

    doc.add_paragraph('Finding 4: Insurer KPI Revealed', style='Intense Quote')
    doc.add_paragraph(
        '"40% of claims volume must be processed fully by AI, without adjuster intervention, '
        'without hiring more people when scaling." This confirms real digitization pressure.'
    )

    doc.add_paragraph('Finding 5: Garages = High Potential, Low Politics', style='Intense Quote')
    doc.add_paragraph(
        'Garages are the most important and powerful party in the claims chain. '
        'They are simple to work with and have clear, quantifiable pain points.'
    )

    add_heading(doc, '6.3 Garage-Specific Insights', 2)
    doc.add_paragraph('From VinFast GMS Gia Lâm workshop manager interview:')
    doc.add_paragraph()
    doc.add_paragraph('• Insurers do NOT accept photos taken by garage')
    doc.add_paragraph('• Adjuster must visit in person (typically once per day for multiple cases)')
    doc.add_paragraph('• Each insurer has own proprietary inspection/photo software')
    doc.add_paragraph('• VAT invoice must be issued BEFORE submitting claim')
    doc.add_paragraph('• Staff manually copy repair services from DMS to MISA accounting software')

    doc.add_page_break()

    # 7. Validated vs Invalidated Hypotheses
    add_heading(doc, '7. VALIDATED VS INVALIDATED HYPOTHESES', 1)

    add_heading(doc, '7.1 Validated Hypotheses', 2)

    add_table(doc,
        ['Hypothesis', 'Evidence', 'Strength'],
        [
            ['Insurers face digitization pressure', 'MIC has 40% AI KPI mandate', 'Strong'],
            ['Garages have real pain', 'Cash flow, DMS-MISA manual copy, slow approvals', 'Strong'],
            ['Market size is significant', '~1,000B VND/year claims, 175K EVs in 2025', 'Strong'],
            ['No good local solution exists', 'FPT.AI = chatbots only, Fermion = old-timer', 'Medium'],
            ['AI Call Center is urgent need', 'Direct request from insurer interview', 'Strong'],
        ]
    )

    add_heading(doc, '7.2 Invalidated Hypotheses', 2)

    add_table(doc,
        ['Hypothesis', 'Reality', 'Implication'],
        [
            ['VinFast wants embedded insurance', 'Too early, not a priority', 'Don\'t depend on OEM'],
            ['Car owners have big claims pain', 'Well-served, any solution is vitamins', 'Don\'t lead with consumer'],
            ['Fraud detection valuable to insurers', '"Luật ngầm" – leadership doesn\'t care', 'Don\'t lead with fraud'],
            ['VinFast will co-develop', 'Won\'t engage until insurers prove it', 'Target insurers first'],
        ]
    )

    add_heading(doc, '7.3 Hypotheses Still to Test', 2)

    add_table(doc,
        ['Hypothesis', 'Proposed Test Method'],
        [
            ['Insurers will pay per claim processed', 'Pricing discussion with pilot insurer (50-100K VND/claim)'],
            ['EV data useful for Risk Score', 'POC with 100-200 claim files + service logs'],
            ['Garages want milestone payment portal', 'Pilot with 5 garages'],
        ]
    )

    doc.add_page_break()

    # 8. Opportunity Sizing
    add_heading(doc, '8. OPPORTUNITY SIZING', 1)

    add_heading(doc, '8.1 Insurance Distribution Revenue Potential', 2)

    add_table(doc,
        ['Parameter', 'Value', 'Notes'],
        [
            ['VinFast annual sales (2025)', '175,099 EVs', 'Actual delivered'],
            ['Digital insurance penetration (baseline)', '8%', 'Industry average'],
            ['Target penetration (with co-branding)', '15%', 'With VinFast marketing partnership'],
            ['Policies at 15% penetration', '26,265 policies/year', 'Calculated'],
            ['Average commission per policy', '1.1 million VND', 'Range: 800K - 1.5M VND'],
            ['Annual revenue potential', '28.9 billion VND', 'At 15% penetration'],
        ]
    )

    add_heading(doc, '8.2 Claims Market Revenue Models', 2)
    doc.add_paragraph('Total motor vehicle claims market: ~1,000 billion VND/year')
    doc.add_paragraph()
    doc.add_paragraph('Potential Revenue Streams:')
    doc.add_paragraph('• Platform fee (pay-per-claim): 50-100K VND per claim processed')
    doc.add_paragraph('• Fast-track premium: Additional fee for expedited 4-hour processing')
    doc.add_paragraph('• Fraud savings share: Percentage of money saved from fraud detection')
    doc.add_paragraph('• Workshop subscription: Fee for garages to join repair quote network')

    add_heading(doc, '8.3 Monthly Run-Rate Potential', 2)
    doc.add_paragraph(
        'Using December 2025 peak (27,649 VinFast vehicles delivered):')
    doc.add_paragraph('• At 15% insurance penetration: 4,147 policies')
    doc.add_paragraph('• At 1.1M VND commission: 4.56 billion VND/month potential')

    doc.add_page_break()

    # 9. Strategic Recommendations
    add_heading(doc, '9. STRATEGIC RECOMMENDATIONS', 1)

    add_heading(doc, '9.1 Target Market', 2)
    doc.add_paragraph(
        'Collision Market = Garages + Insurers (not car owners)')
    doc.add_paragraph()
    doc.add_paragraph(
        'This represents a strategic pivot from the original consumer-focused embedded insurance concept. '
        'The Collision Market approach aligns with validated pain points and offers clearer path to revenue.'
    )

    add_heading(doc, '9.2 Recommended Entry Points (Prioritized)', 2)

    doc.add_paragraph('Priority 1: AI Call Center')
    doc.add_paragraph('• Explicit demand from insurer interviews')
    doc.add_paragraph('• Quick win for insurer leadership')
    doc.add_paragraph('• Lower technical risk than full claims automation')
    doc.add_paragraph('• Concrete, demonstrable value')
    doc.add_paragraph()

    doc.add_paragraph('Priority 2: Garage Portal')
    doc.add_paragraph('• High pain, low politics')
    doc.add_paragraph('• Clear ROI (cash flow improvement)')
    doc.add_paragraph('• Garages are most powerful party in claims chain')
    doc.add_paragraph('• Simpler stakeholder management')
    doc.add_paragraph()

    doc.add_paragraph('Priority 3: Claims Triage System')
    doc.add_paragraph('• Meets insurer 40% AI processing KPI')
    doc.add_paragraph('• Enables auto-approval for low-risk claims')
    doc.add_paragraph('• Focuses human adjusters on complex cases')

    add_heading(doc, '9.3 NOT Recommended to Lead With', 2)
    doc.add_paragraph('• Fraud detection – NOT a priority for insurer leadership')
    doc.add_paragraph('• Consumer claims experience – vitamins, not painkillers')
    doc.add_paragraph('• VinFast integration – they won\'t co-develop until insurers prove it')

    add_heading(doc, '9.4 Go-to-Market Strategy', 2)
    doc.add_paragraph('Target: Insurers first, not OEMs')
    doc.add_paragraph()
    doc.add_paragraph('Tactics:')
    doc.add_paragraph('• Pitch multiple insurers simultaneously to create competitive pressure')
    doc.add_paragraph('• Focus messaging on: innovation, operational efficiency, fast measurable results')
    doc.add_paragraph('• Emphasize personal wins for leadership (recognition, bonuses, shareholder credibility)')
    doc.add_paragraph('• Position as early partnership for light POC, not major commitment')

    add_heading(doc, '9.5 POC Approach', 2)
    doc.add_paragraph('"Solve a smaller problem first to validate desirability risk."')
    doc.add_paragraph()
    doc.add_paragraph('Recommended POC scope:')
    doc.add_paragraph('• 5 garages (mix of VinFast authorized + private)')
    doc.add_paragraph('• 1 strong insurer partner')
    doc.add_paragraph('• Single entry point (AI Call Center OR Garage Portal)')
    doc.add_paragraph('• 2-3 month validation period')

    doc.add_page_break()

    # 10. Implementation Roadmap
    add_heading(doc, '10. IMPLEMENTATION ROADMAP', 1)

    add_heading(doc, '10.1 Phase 1: Pilot (2025)', 2)
    doc.add_paragraph('Duration: 3-6 months')
    doc.add_paragraph()
    doc.add_paragraph('Objectives:')
    doc.add_paragraph('• Secure 1 insurer partner willing to pilot')
    doc.add_paragraph('• Onboard 5 target garages (VinFast authorized + private)')
    doc.add_paragraph('• Define MVP scope for chosen entry point')
    doc.add_paragraph('• Build narrow MVP')
    doc.add_paragraph('• Prove 4-hour/48-hour claims commitment')
    doc.add_paragraph()
    doc.add_paragraph('Success Metrics:')
    doc.add_paragraph('• Claims processing time reduction')
    doc.add_paragraph('• User satisfaction scores')
    doc.add_paragraph('• System uptime and reliability')

    add_heading(doc, '10.2 Phase 2: Expand (2026)', 2)
    doc.add_paragraph('Duration: 12 months')
    doc.add_paragraph()
    doc.add_paragraph('Objectives:')
    doc.add_paragraph('• Expand to all cars (not just VinFast)')
    doc.add_paragraph('• Build centralized data center with 3+ pioneer insurers')
    doc.add_paragraph('• Add second entry point (triage or call center)')
    doc.add_paragraph('• Develop standardized APIs for insurer integration')
    doc.add_paragraph()
    doc.add_paragraph('Success Metrics:')
    doc.add_paragraph('• Number of insurers integrated')
    doc.add_paragraph('• Claims volume processed')
    doc.add_paragraph('• Revenue from platform fees')

    add_heading(doc, '10.3 Phase 3: Scale (2027)', 2)
    doc.add_paragraph('Duration: Ongoing')
    doc.add_paragraph()
    doc.add_paragraph('Objectives:')
    doc.add_paragraph('• Nationwide coverage for motor vehicles')
    doc.add_paragraph('• Include motorcycles in platform')
    doc.add_paragraph('• Establish national claims database')
    doc.add_paragraph('• Achieve industry standard status')
    doc.add_paragraph()
    doc.add_paragraph('Vision:')
    doc.add_paragraph(
        'Electrik becomes the national claims data platform for motor vehicles in Vietnam, '
        'aligned with government digital transformation initiatives and serving as the '
        'infrastructure layer for the entire insurance industry.'
    )

    doc.add_page_break()

    # 11. Appendices
    add_heading(doc, '11. APPENDICES', 1)

    add_heading(doc, '11.1 Research Documents', 2)

    add_table(doc,
        ['#', 'Document', 'Date', 'Content'],
        [
            ['01', 'problem-statement.md', 'Jul 31, 2025', 'BeeVN gaps, opportunities'],
            ['02', 'value-proposition.md', 'Aug 2, 2025', 'Embedded insurance VP + landing page'],
            ['03', 'embedded-insurance-research.md', 'Aug 6, 2025', 'Market context'],
            ['04', 'market-research.md', 'Aug 7, 2025', 'VinFast sales, revenue sizing'],
            ['05', 'pivot-to-claims.md', 'Aug 22-24, 2025', 'Pivot, competitive landscape'],
            ['06', 'claims-market-research.md', 'Aug 26, 2025', 'Market size, fraud types'],
            ['07', 'cto-pitch-mvp.md', 'Aug 29, 2025', 'MVP roadmap'],
            ['08', 'federated-learning-research.md', 'Sep 22, 2025', 'Privacy-preserving ML'],
            ['09', 'claim-triage-system.md', 'Sep 25, 2025', 'Triage concept'],
            ['10', 'customer-interviews.md', 'Sep 24-25, 2025', 'Critical evidence'],
            ['11', 'garage-interview.md', 'Sep 25, 2025', 'Workshop manager insights'],
            ['12', 'marketing-copy.md', 'Oct 10, 2025', 'Positioning by audience'],
            ['13', 'product-concept-sep5.md', 'Sep 5, 2025', 'Historical product concept'],
            ['14', 'pivot-vpc-bmc.md', 'Aug 22, 2025', 'Full VPC/BMC'],
            ['15', 'weekly-validation-sep30.md', 'Sep 30, 2025', 'Synthesis + start small'],
            ['16', 'formal-proposal.md', 'Sep 15, 2025', 'Đề án for stakeholders'],
            ['17', 'competitive-analysis.md', 'Sep 18, 2025', '12 solutions analyzed'],
        ]
    )

    add_heading(doc, '11.2 Key Stakeholders', 2)

    add_table(doc,
        ['Stakeholder', 'Role', 'Relationship'],
        [
            ['BeeVN', 'Vingroup in-house broker', 'Primary client'],
            ['VinFast', 'EV manufacturer', 'Ecosystem partner'],
            ['GSM/XanhSM', 'Taxi/ride-hailing service', 'Distribution channel'],
            ['PVI', 'Insurance partner', 'Existing BeeVN partner'],
            ['Bảo Việt', 'Major insurer', 'Potential partner'],
            ['BIC', 'Insurer', 'VinFast partnership (Dec 2024)'],
            ['VBI', 'Insurer', 'VinFast partnership (Dec 2024)'],
            ['PTI', 'Insurer', 'VinFast partnership (Dec 2024)'],
            ['BSH', 'Insurer', 'VinFast partnership (Dec 2024)'],
            ['VNI/DBV', 'Insurer', 'VinFast partnership (Dec 2024)'],
        ]
    )

    add_heading(doc, '11.3 Regulatory References', 2)
    doc.add_paragraph('• Luật Kinh doanh bảo hiểm (Insurance Business Law)')
    doc.add_paragraph('• Nghị định 67/2023/NĐ-CP: Mandatory liability insurance requirements')
    doc.add_paragraph('• Politburo Resolution 57-NQ/TW: Science, technology, innovation, digital transformation')
    doc.add_paragraph('• Politburo Resolution 68-NQ/TW: Private sector development')
    doc.add_paragraph('• December 2024: VinFast agreement with 7 insurers for expedited claims')

    add_heading(doc, '11.4 Glossary', 2)

    add_table(doc,
        ['Term', 'Definition'],
        [
            ['BeeVN', 'Vingroup\'s in-house insurance broker'],
            ['Collision Market', 'B2B market of garages and insurers handling vehicle repairs'],
            ['DMS', 'Dealer Management System (OEM-provided software)'],
            ['MISA', 'Popular Vietnamese accounting software'],
            ['Luật ngầm', '"Unwritten rule" - referring to accepted industry practices'],
            ['VPC', 'Value Proposition Canvas'],
            ['BMC', 'Business Model Canvas'],
            ['POC', 'Proof of Concept'],
            ['TPA', 'Third Party Administrator'],
            ['UBI', 'Usage-Based Insurance'],
            ['CQLTT', 'Cục Quản lý thị trường (Market Management Authority)'],
        ]
    )

    # Save document
    output_path = '/Users/nntrvy/osr-company/electrik-project/reports/electrik-market-research-report.docx'
    doc.save(output_path)
    print(f'Report saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()
