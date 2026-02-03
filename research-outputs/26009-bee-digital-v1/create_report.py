#!/usr/bin/env python3
"""
Generate comprehensive market research report for BeeVN/Electrik project
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add a heading with consistent styling"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    doc.add_paragraph()
    return table

def create_report():
    doc = Document()

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title = doc.add_heading('ELECTRIK', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Comprehensive Market Research Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle2 = doc.add_paragraph('Insurance Technology Platform for Vietnam\'s Collision Market')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Prepared for: BeeVN (Vingroup Insurance Broker)\n').bold = True
    info.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}\n')
    info.add_run('Prepared by: OS Research Company\n')
    info.add_run('Research Period: July - October 2025')

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    add_heading(doc, 'Table of Contents', 1)

    toc_items = [
        '1. Executive Summary',
        '2. Market Overview',
        '   2.1 Vietnam Auto Insurance Market',
        '   2.2 Electric Vehicle Market',
        '   2.3 Claims Market Size',
        '3. Current State Analysis',
        '   3.1 BeeVN\'s Current Situation',
        '   3.2 Industry Pain Points',
        '4. Stakeholder Analysis',
        '   4.1 Car Owners',
        '   4.2 Insurance Companies',
        '   4.3 Repair Garages',
        '   4.4 VinFast/OEM',
        '5. Competitive Landscape',
        '   5.1 Local Players',
        '   5.2 International References',
        '   5.3 Blue Ocean Positioning',
        '6. Customer Interview Insights',
        '   6.1 Interview Summary',
        '   6.2 Key Findings',
        '   6.3 Garage-Specific Insights',
        '7. Hypothesis Validation Results',
        '8. Opportunity Sizing',
        '9. Strategic Recommendations',
        '10. Implementation Roadmap',
        '11. Appendices',
    ]

    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # =========================================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================================
    add_heading(doc, '1. Executive Summary', 1)

    doc.add_paragraph(
        'This report presents comprehensive market research conducted over a 3-month period '
        '(July - October 2025) to validate an insurance technology opportunity within the '
        'Vingroup ecosystem. The research involved extensive stakeholder interviews, competitive '
        'analysis, and market sizing to identify the most viable entry point for BeeVN\'s '
        'digital transformation.'
    )

    add_heading(doc, 'Project Evolution', 2)
    doc.add_paragraph(
        'The project began as an "embedded insurance" concept—integrating insurance products '
        'directly into Vingroup\'s digital apps (VinFast, GSM, etc.). However, through rigorous '
        'validation, this hypothesis was refuted. Vingroup indicated it was "not the right time" '
        'and the initiative was not a priority.'
    )
    doc.add_paragraph(
        'The project pivoted to focus on claims processing after Vingroup confirmed pain points '
        'in this area. Further customer interviews revealed that the real opportunity lies in '
        'serving the Collision Market—specifically garages and insurers—rather than end consumers.'
    )

    add_heading(doc, 'Key Findings', 2)

    findings = [
        ['Finding', 'Implication'],
        ['Car owners in Vietnam are relatively well-served', 'Build "vitamins" not "painkillers" for consumers'],
        ['Fraud detection is NOT a priority for insurer leadership', 'Do not lead with fraud detection messaging'],
        ['AI Call Center has explicit urgent demand', 'Best entry point for quick wins'],
        ['Garages are the most powerful party in claims chain', 'Focus on garage pain points first'],
        ['Insurers have 40% AI processing KPI mandate', 'Real digitization pressure exists'],
        ['VinFast will not co-develop until insurers prove value', 'Target insurers first, not OEMs'],
    ]
    add_table(doc, findings[0], findings[1:])

    add_heading(doc, 'Market Opportunity', 2)
    doc.add_paragraph(
        'The Vietnam motor vehicle claims market represents approximately 1,000 billion VND per year. '
        'With VinFast delivering 175,099 EVs in 2025 (a record), the EV-specific insurance market '
        'is growing rapidly. At 15% digital insurance penetration with average commission of 1.1M VND '
        'per policy, the distribution opportunity alone is approximately 29 billion VND annually.'
    )

    add_heading(doc, 'Recommended Direction', 2)
    doc.add_paragraph().add_run('Target Market: ').bold = True
    doc.paragraphs[-1].add_run('Collision Market (Garages + Insurers), not car owners')

    doc.add_paragraph().add_run('Priority Entry Points:').bold = True
    entry_points = doc.add_paragraph()
    entry_points.add_run('\n1. AI Call Center – Explicit demand from insurer interviews')
    entry_points.add_run('\n2. Garage Portal – High pain, low politics, clear ROI')
    entry_points.add_run('\n3. Claims Triage System – Meets insurer KPI requirements')

    doc.add_paragraph().add_run('POC Approach: ').bold = True
    doc.paragraphs[-1].add_run('Start with 5 garages + 1 strong insurer partner before scaling')

    doc.add_page_break()

    # =========================================================================
    # 2. MARKET OVERVIEW
    # =========================================================================
    add_heading(doc, '2. Market Overview', 1)

    add_heading(doc, '2.1 Vietnam Auto Insurance Market', 2)

    doc.add_paragraph(
        'Vietnam\'s automotive market has shown strong growth, with total vehicle sales reaching '
        'approximately 494,000 units in 2024. The insurance industry is undergoing digital '
        'transformation pressure, though adoption of advanced technologies remains limited.'
    )

    auto_market_data = [
        ['Metric', '2024 Value', 'Notes'],
        ['Total cars sold', '~494,000', 'VAMA statistics'],
        ['EVs sold', '87,000+', '18% market share'],
        ['VinFast EV market share', '90%+', 'Dominant player'],
        ['Motor claims payout', '~1,000B VND/year', 'Industry total'],
        ['Estimated fraud rate', '~15%', 'Of total claims costs'],
    ]
    add_table(doc, auto_market_data[0], auto_market_data[1:])

    add_heading(doc, '2.2 Electric Vehicle Market', 2)

    doc.add_paragraph(
        'Vietnam\'s EV market is experiencing explosive growth, driven primarily by VinFast. '
        'The company achieved record deliveries in 2025, nearly doubling its 2024 performance.'
    )

    ev_market_data = [
        ['Year/Metric', 'Value', 'Source'],
        ['VinFast 2025 deliveries', '175,099 EVs', 'VinFast official'],
        ['December 2025 (peak month)', '27,649 units', 'Monthly record'],
        ['Vietnam EV market 2025', '$2.93 billion', 'Mordor Intelligence'],
        ['Vietnam EV market 2030 (projected)', '$6.69 billion', 'Mordor Intelligence'],
        ['CAGR 2025-2030', '~18%', 'Mordor Intelligence'],
        ['EVs on road by 2028 (projected)', '1 million', 'VAMA'],
    ]
    add_table(doc, ev_market_data[0], ev_market_data[1:])

    add_heading(doc, 'VinFast 2025 Model Breakdown', 3)

    model_data = [
        ['Model', 'Units Sold', 'Segment'],
        ['VF 3', '44,585', 'Entry-level'],
        ['VF 5', '43,913', 'Compact'],
        ['Limo Green', '27,127', 'B2B/Taxi'],
        ['VF 6', '23,291', 'Mid-range'],
        ['Herio Green', '12,568', 'B2B'],
        ['VF 7', '9,653', 'Premium'],
    ]
    add_table(doc, model_data[0], model_data[1:])

    add_heading(doc, '2.3 Claims Market Size', 2)

    doc.add_paragraph(
        'The motor vehicle claims market represents a significant opportunity. Major insurers '
        'process hundreds of billions in claims annually, with substantial inefficiencies in '
        'the current process.'
    )

    claims_data = [
        ['Insurer', 'Annual Claims (Approx.)', 'Notes'],
        ['PTI', '1,200-1,400B VND', 'Peak in 2022-2023'],
        ['MIC', '800-1,000B VND', 'Digital transformation focus'],
        ['BSH', '800-1,000B VND', 'VinFast partner'],
        ['PJICO', '500-700B VND', 'Traditional player'],
        ['DBV', '200-400B VND', 'Growing segment'],
    ]
    add_table(doc, claims_data[0], claims_data[1:])

    doc.add_page_break()

    # =========================================================================
    # 3. CURRENT STATE ANALYSIS
    # =========================================================================
    add_heading(doc, '3. Current State Analysis', 1)

    add_heading(doc, '3.1 BeeVN\'s Current Situation', 2)

    doc.add_paragraph(
        'BeeVN operates as Vingroup\'s in-house insurance broker, managing commercial insurance '
        'for all subsidiaries and negotiating claims with insurers. Their current technology '
        'infrastructure is limited, creating dependencies on external partners.'
    )

    add_heading(doc, 'Current Partnerships', 3)
    partnerships = [
        ['Partner', 'Product', 'Market'],
        ['PVI Digital', 'Trip insurance for GSM passengers', 'Vietnam'],
        ['Malayan', 'Driver savings insurance', 'GSM Philippines'],
    ]
    add_table(doc, partnerships[0], partnerships[1:])

    add_heading(doc, 'GSM Trip Insurance Case Study', 3)
    doc.add_paragraph(
        'BeeVN partnered with PVI Digital to sell trip insurance to GSM ride-hailing passengers '
        'at 2,000 VND per trip. Initial results were strong due to automatic opt-in:'
    )

    gsm_data = [
        ['Period', 'Monthly Revenue', 'Change'],
        ['Before regulation', '15B VND/month', 'Baseline'],
        ['After CQLTT regulation', '8B VND/month', '-47%'],
    ]
    add_table(doc, gsm_data[0], gsm_data[1:])

    doc.add_paragraph(
        'The sharp decline occurred when market regulators (CQLTT) required explicit opt-in '
        'consent from customers. This revealed BeeVN\'s core vulnerability: lack of control '
        'over systems, integration, and customer demand.'
    )

    add_heading(doc, '3.2 Industry Pain Points', 2)

    pain_points = [
        ['Pain Point', 'Description', 'Impact'],
        ['Slow claims process', '15-30 days typical resolution', 'Customer frustration'],
        ['Manual paperwork', 'Error-prone, labor-intensive', 'Operational cost'],
        ['No unified system', 'Customer-garage-insurer disconnected', 'Inefficiency'],
        ['No EV-specific solution', 'Traditional systems don\'t handle EV parts/battery', 'Market gap'],
        ['Fraud losses', '~15% of claims costs', 'Financial leakage'],
        ['Declining trust', 'Public perception of insurance is negative', 'Market size limited'],
    ]
    add_table(doc, pain_points[0], pain_points[1:])

    add_heading(doc, 'Five Common Insurance Fraud Types in Vietnam', 3)
    fraud_types = [
        ['Type', 'Description'],
        ['1. Falsify loss timing', 'Claim for damage that occurred before policy purchase'],
        ['2. Double-dip claims', 'Already compensated by third party, create fake scene for insurer'],
        ['3. Hide violations', 'No license or under influence, lie to receive compensation'],
        ['4. Inflate loss value', 'Replace old/broken parts, stage scene to legitimize'],
        ['5. Cross-border fraud', 'Damage outside Vietnam, drag car back, report fake location'],
    ]
    add_table(doc, fraud_types[0], fraud_types[1:])

    doc.add_page_break()

    # =========================================================================
    # 4. STAKEHOLDER ANALYSIS
    # =========================================================================
    add_heading(doc, '4. Stakeholder Analysis', 1)

    doc.add_paragraph(
        'Understanding the pain levels, power dynamics, and priorities of each stakeholder is '
        'critical for product-market fit. Our research revealed unexpected findings about where '
        'the real opportunities lie.'
    )

    add_heading(doc, '4.1 Car Owners', 2)

    doc.add_paragraph().add_run('Pain Level: LOW').bold = True
    doc.add_paragraph(
        'Contrary to initial assumptions, Vietnamese car owners are relatively well-served '
        'compared to foreign markets. Insurance claims, while not perfect, do not represent '
        'a major pain point for most consumers.'
    )

    doc.add_paragraph().add_run('Key Insight: ').bold = True
    doc.paragraphs[-1].add_run(
        '"Any solution targeting car owners directly would be vitamins, not painkillers. '
        'Focus on making great vitamins, not great painkillers."'
    )

    doc.add_paragraph().add_run('Implication: ').bold = True
    doc.paragraphs[-1].add_run('Do not build a consumer-facing claims app as the primary product.')

    add_heading(doc, '4.2 Insurance Companies', 2)

    doc.add_paragraph().add_run('Pain Level: MEDIUM (Real but not urgent)').bold = True

    doc.add_paragraph(
        'Insurers face digitization pressure from leadership, but current solutions are '
        '"good enough" for most. The pain is real and significant, but not urgent enough '
        'to drive immediate action without clear ROI.'
    )

    add_heading(doc, 'Key Findings from MIC Interview', 3)
    mic_findings = [
        ['Finding', 'Evidence Strength'],
        ['40% of claims must be processed by AI without human intervention', 'Strong - specific KPI'],
        ['Digitization mandate from leadership exists', 'Medium - general directive'],
        ['Fraud detection is NOT a priority', 'Strong - "luật ngầm" (unwritten rule)'],
        ['AI Call Center is an urgent need', 'Strong - explicit request'],
        ['Current internal solutions are "good enough"', 'Medium - no urgency to change'],
    ]
    add_table(doc, mic_findings[0], mic_findings[1:])

    doc.add_paragraph().add_run('Critical Insight on Fraud Detection: ').bold = True
    doc.paragraphs[-1].add_run(
        'Fraud at some level is considered an "unwritten rule" (luật ngầm) in the industry. '
        'Leadership does not prioritize fraud detection because it is very difficult to '
        'measure and quantify actual savings. Do NOT lead with fraud detection as a value proposition.'
    )

    add_heading(doc, '4.3 Repair Garages', 2)

    doc.add_paragraph().add_run('Pain Level: HIGH').bold = True
    doc.add_paragraph().add_run('Power Level: HIGHEST in the claims chain').bold = True

    doc.add_paragraph(
        'Garages emerged as the most promising stakeholder segment. They have real, measurable '
        'pain points and are the most powerful party in the claims ecosystem. They are also '
        'simpler to work with—less political than insurers or OEMs.'
    )

    garage_data = [
        ['Metric', 'Value', 'Notes'],
        ['Average monthly revenue', '3.6B VND', 'Per garage'],
        ['Insurance customer share', '~15%', 'Of total revenue'],
        ['Insurance revenue', '~600M VND/month', 'Per garage'],
        ['Insurance work margin', 'Higher than warranty', 'Includes labor component'],
    ]
    add_table(doc, garage_data[0], garage_data[1:])

    add_heading(doc, 'Garage Pain Points', 3)
    garage_pains = [
        ['Pain Point', 'Current Workaround', 'Opportunity'],
        ['Document preparation time-consuming', 'Manual collection', 'Automation'],
        ['Insurers don\'t notify on incomplete docs', 'Garage must chase', 'Status tracking'],
        ['DMS doesn\'t integrate with MISA accounting', 'Manual copy for every invoice', 'Integration'],
        ['Insurers don\'t accept garage photos', 'Adjuster must visit in person', 'Verified photo system'],
        ['Slow insurance approval', 'Wait and follow up', 'Faster processing'],
        ['Unstable cash flow', 'Advance repair costs', 'Milestone payments'],
    ]
    add_table(doc, garage_pains[0], garage_pains[1:])

    add_heading(doc, '4.4 VinFast/OEM', 2)

    doc.add_paragraph().add_run('Pain Level: LOW (for insurance specifically)').bold = True

    doc.add_paragraph(
        'VinFast\'s primary focus is on vehicle sales and their core ecosystem (GSM taxi service). '
        'Insurance experience improvement is a low priority in their strategic roadmap.'
    )

    vinfast_findings = [
        ['Finding', 'Implication'],
        ['Would consider integration IF insurers have proven solution', 'Insurers must prove value first'],
        ['Will NOT co-develop from scratch', 'Don\'t spend resources trying to partner early'],
        ['Better insurance experience is small priority', 'Focus elsewhere'],
        ['GSM is their main customer concern', 'B2B fleet might be entry point later'],
        ['Everything requires Vingroup chairman approval', 'Long decision cycles'],
    ]
    add_table(doc, vinfast_findings[0], vinfast_findings[1:])

    doc.add_page_break()

    # =========================================================================
    # 5. COMPETITIVE LANDSCAPE
    # =========================================================================
    add_heading(doc, '5. Competitive Landscape', 1)

    add_heading(doc, '5.1 Local Players', 2)

    local_players = [
        ['Player', 'Strength', 'Weakness', 'Opportunity for Electrik'],
        ['Fermion/Merimen', 'Large system, insurer relationships', 'Old-timer, not EV-optimized', 'Modern EV-first approach'],
        ['FPT.AI', 'Strong AI (OCR, NLP, Car Damage)', 'Mostly chatbots, not end-to-end', 'Full claims workflow'],
        ['AiCycle.ai', 'Local, AI-focused', 'Pivoting to distribution', 'Fill the tech gap they\'re leaving'],
        ['Bảo Việt/PVI/PTI apps', 'Existing customer base', 'Basic digitization only, no AI', 'Superior CX + AI'],
        ['Global Care', 'Platform expertise (acquired by Huize)', 'Now Chinese-owned', 'Independent Vietnamese solution'],
    ]
    add_table(doc, local_players[0], local_players[1:])

    doc.add_paragraph().add_run('Note on Fermion/Merimen: ').bold = True
    doc.paragraphs[-1].add_run(
        'They are NOT a direct competitor. They are an old-timer infrastructure provider '
        'serving insurers. Electrik targets a different position—customer-centric, EV-specialized.'
    )

    add_heading(doc, '5.2 International References', 2)

    intl_players = [
        ['Company', 'Country', 'Key Strength', 'What to Learn'],
        ['Lemonade', 'US', 'Best-in-class CX', '"Wow" experience, trust through transparency'],
        ['FRISS', 'Netherlands', 'Fraud scoring AI', 'Flip to "Honest Score" (positive framing)'],
        ['Ravin.ai', 'Israel', 'AI damage detection', 'One-touch video claims'],
        ['CCC', 'US', 'Collision market platform', 'Reference model for garage-insurer platform'],
        ['Root Insurance', 'US', 'Usage-based insurance (UBI)', 'Light UBI version for EV'],
        ['PasarPolis', 'Indonesia', 'Ecosystem integration', 'Plug into existing apps, don\'t go solo'],
        ['Igloo', 'Singapore', 'API-first B2B2C', 'Fast partner integration model'],
    ]
    add_table(doc, intl_players[0], intl_players[1:])

    add_heading(doc, '"Honest Score" Concept', 3)
    doc.add_paragraph(
        'Instead of "Fraud Score" (negative, suspicious framing), Electrik should use '
        '"Honest Score"—same AI technology, but positive framing that rewards transparent '
        'customers with faster processing. This builds trust instead of suspicion.'
    )

    add_heading(doc, '5.3 Blue Ocean Positioning', 2)

    doc.add_paragraph(
        'The competitive landscape can be mapped on two axes: Customer Experience and Automation. '
        'Electrik\'s opportunity is in the upper-right quadrant—high on both dimensions, '
        'specifically optimized for the EV ecosystem.'
    )

    positioning = [
        ['Player', 'Automation', 'Customer Experience', 'Position'],
        ['Electrik (target)', 'High', 'High', 'Blue Ocean - EV-first'],
        ['Fermion/Merimen', 'High', 'Medium', 'Insurer-centric'],
        ['FPT.AI', 'Medium-High', 'Low', 'Backend only'],
        ['Insurtech apps (Saladin, OPES)', 'Low-Medium', 'High', 'Consumer front-end'],
        ['Traditional adjusters', 'Low', 'Low', 'Legacy'],
        ['Local insurer apps', 'Low', 'Medium', 'Basic digitization'],
    ]
    add_table(doc, positioning[0], positioning[1:])

    doc.add_page_break()

    # =========================================================================
    # 6. CUSTOMER INTERVIEW INSIGHTS
    # =========================================================================
    add_heading(doc, '6. Customer Interview Insights', 1)

    doc.add_paragraph(
        'Customer interviews conducted in September 2025 provided the strongest evidence for '
        'strategic decisions. These findings should be weighted heavily in any future planning.'
    )

    add_heading(doc, '6.1 Interview Summary', 2)

    interviews = [
        ['Date', 'Interviewee', 'Role', 'Organization'],
        ['Sep 24, 2025', 'Anh Phong', 'Head of Claims', 'MIC (insurer)'],
        ['Sep 24, 2025', 'Anh Bảo', 'Product Manager', 'XanhSM/VinFast'],
        ['Sep 25, 2025', 'Workshop Manager', 'Operations', 'VinFast GMS Gia Lâm'],
        ['Sep 24-25', 'Claims Team', 'Claims Processing', 'PJICO'],
    ]
    add_table(doc, interviews[0], interviews[1:])

    add_heading(doc, '6.2 Key Findings', 2)

    add_heading(doc, 'Finding 1: Digitization Pressure is Real', 3)
    doc.add_paragraph().add_run('Evidence Strength: Medium').bold = True
    doc.add_paragraph(
        'Claims departments are under pressure to digitize and apply AI. The barrier is not '
        'willingness but capability—they lack resources to build in-house and haven\'t found '
        'suitable Vietnam solutions that integrate with existing systems.'
    )
    doc.add_paragraph().add_run('Specific KPI: ').bold = True
    doc.paragraphs[-1].add_run(
        '"40% of claims volume must be processed fully by AI, without adjuster intervention, '
        'without hiring more people when scaling."'
    )

    add_heading(doc, 'Finding 2: Fraud Detection is NOT a Priority', 3)
    doc.add_paragraph().add_run('Evidence Strength: Strong').bold = True
    doc.add_paragraph(
        'Contrary to assumptions, fraud detection is not a major concern for insurer leadership. '
        'Fraud at some level is considered an "unwritten rule" (luật ngầm). The difficulty in '
        'measuring and quantifying savings makes it hard to justify investment.'
    )
    doc.add_paragraph().add_run('Implication: ').bold = True
    doc.paragraphs[-1].add_run('Do NOT position Electrik around fraud detection as primary value.')

    add_heading(doc, 'Finding 3: AI Call Center = Urgent Need', 3)
    doc.add_paragraph().add_run('Evidence Strength: Strong').bold = True
    doc.add_paragraph(
        'Some insurers have researched AI call centers but found no effective solution. '
        'During interviews, a partner immediately asked: "Can you build this solution for me?" '
        'This explicit demand signal indicates AI Call Center is the fastest entry point.'
    )

    add_heading(doc, 'Finding 4: VinFast Won\'t Co-Develop', 3)
    doc.add_paragraph().add_run('Evidence Strength: Weak-Medium').bold = True
    doc.add_paragraph(
        'VinFast would consider integration if insurers have a proven, reliable solution. '
        'However, they will not co-develop from scratch. Better insurance experience is a '
        'small priority—they primarily care about GSM as their main customer.'
    )

    add_heading(doc, 'Finding 5: Garages = High Potential, Low Politics', 3)
    doc.add_paragraph().add_run('Evidence Strength: Strong').bold = True
    doc.add_paragraph(
        'Garages emerged as the most promising segment. They are the most powerful party '
        'in the claims chain, have clear pain points, and are simpler to work with than '
        'insurers or OEMs. The 151 VinFast workshops plus private garage network offer scale.'
    )

    add_heading(doc, '6.3 Garage-Specific Insights', 2)

    doc.add_paragraph(
        'The interview with VinFast GMS Gia Lâm workshop manager revealed specific operational '
        'pain points and revenue data.'
    )

    add_heading(doc, 'Current System Issues', 3)
    system_issues = [
        ['System', 'Purpose', 'Problem'],
        ['DMS (OEM-provided)', 'Manage processes, parts ordering', 'Clunky, hard to use'],
        ['MISA (accounting)', 'Issue invoices', 'No integration with DMS'],
    ]
    add_table(doc, system_issues[0], system_issues[1:])

    doc.add_paragraph(
        'For insurance claims, VAT invoice must be issued BEFORE submitting claim. '
        'Staff must manually copy repair services from DMS to MISA for every case.'
    )

    add_heading(doc, 'Photo Documentation Process', 3)
    doc.add_paragraph().add_run('Critical finding: ').bold = True
    doc.paragraphs[-1].add_run(
        'Insurers do NOT accept photos taken by garage. An adjuster must visit in person '
        'to take photos using their proprietary inspection apps. Adjusters typically visit '
        'once per day to handle multiple cases.'
    )
    doc.add_paragraph().add_run('Opportunity: ').bold = True
    doc.paragraphs[-1].add_run('Verified photo system could eliminate adjuster visits for simple cases.')

    doc.add_page_break()

    # =========================================================================
    # 7. HYPOTHESIS VALIDATION RESULTS
    # =========================================================================
    add_heading(doc, '7. Hypothesis Validation Results', 1)

    doc.add_paragraph(
        'The validation process tested multiple hypotheses about market desirability, '
        'feasibility, and viability. Results are categorized below.'
    )

    add_heading(doc, 'Validated Hypotheses ✓', 2)

    validated = [
        ['Hypothesis', 'Evidence', 'Confidence'],
        ['Insurers face digitization pressure', 'MIC has 40% AI KPI mandate', 'High'],
        ['Garages have real pain points', 'Cash flow, DMS-MISA gap, slow approvals', 'High'],
        ['Market size is significant', '~1,000B VND/year claims, 175K EVs in 2025', 'High'],
        ['No good local solution exists', 'FPT.AI = chatbots, Fermion = old, AiCycle pivoting', 'High'],
        ['AI Call Center is urgent need', 'Direct explicit request from insurer', 'High'],
    ]
    add_table(doc, validated[0], validated[1:])

    add_heading(doc, 'Invalidated Hypotheses ✗', 2)

    invalidated = [
        ['Hypothesis', 'Reality', 'Learning'],
        ['VinFast wants embedded insurance', 'Too early, not a priority', 'Don\'t depend on OEM'],
        ['Car owners have big claims pain', 'Well-served compared to global standards', 'Build vitamins'],
        ['Fraud detection is valuable to insurers', '"Luật ngầm" - leadership doesn\'t care', 'Don\'t lead with this'],
        ['VinFast will co-develop', 'Won\'t engage until insurers prove it', 'Target insurers first'],
    ]
    add_table(doc, invalidated[0], invalidated[1:])

    add_heading(doc, 'Hypotheses Still to Test', 2)

    to_test = [
        ['Hypothesis', 'Test Method', 'Priority'],
        ['Insurers will pay per claim processed', 'Pricing discussion with pilot insurer', 'High'],
        ['EV data useful for Risk Score', 'POC with 100-200 claim files', 'Medium'],
        ['Garages want milestone payment portal', 'Pilot with 5 garages', 'High'],
    ]
    add_table(doc, to_test[0], to_test[1:])

    doc.add_page_break()

    # =========================================================================
    # 8. OPPORTUNITY SIZING
    # =========================================================================
    add_heading(doc, '8. Opportunity Sizing', 1)

    add_heading(doc, '8.1 Insurance Distribution Opportunity', 2)

    doc.add_paragraph(
        'If Electrik enables digital insurance distribution through the VinFast ecosystem:'
    )

    distribution_calc = [
        ['Metric', 'Value', 'Source'],
        ['VinFast annual sales', '175,099 EVs', '2025 actual'],
        ['Target digital penetration', '15%', 'With co-branding (vs 8% baseline)'],
        ['Policies sold', '26,265/year', 'Calculated'],
        ['Average commission', '1.1M VND', 'Range: 800K - 1.5M'],
        ['Annual revenue potential', '~29B VND', 'Calculated'],
    ]
    add_table(doc, distribution_calc[0], distribution_calc[1:])

    add_heading(doc, '8.2 Claims Processing Opportunity', 2)

    doc.add_paragraph('The claims market offers multiple revenue model possibilities:')

    claims_revenue = [
        ['Revenue Model', 'Description', 'Potential'],
        ['Platform fee (PAYG)', 'Fee per claim processed', '50-100K VND/claim'],
        ['Fast-track premium', 'Premium for expedited processing', 'Additional 20-30%'],
        ['Fraud savings share', '% of money saved from detection', '5-10% of savings'],
        ['Garage subscription', 'Fee to join repair quote network', 'Monthly SaaS'],
    ]
    add_table(doc, claims_revenue[0], claims_revenue[1:])

    add_heading(doc, '8.3 Market Size Context', 2)

    market_context = [
        ['Market', 'Size', 'Electrik Target'],
        ['Total motor claims (annual)', '~1,000B VND', '1-5% initial'],
        ['VinFast insurance revenue', '~29B VND potential', '15% penetration'],
        ['Garage insurance revenue', '~600M VND/garage/month', '151 VF garages + private'],
        ['Global embedded insurance (2029)', '$703B', 'Long-term reference'],
    ]
    add_table(doc, market_context[0], market_context[1:])

    doc.add_page_break()

    # =========================================================================
    # 9. STRATEGIC RECOMMENDATIONS
    # =========================================================================
    add_heading(doc, '9. Strategic Recommendations', 1)

    add_heading(doc, '9.1 Target Market', 2)

    doc.add_paragraph().add_run('Primary Target: Collision Market').bold = True
    doc.add_paragraph('This includes:')
    doc.add_paragraph('• Repair garages (VinFast authorized + private network)')
    doc.add_paragraph('• Insurance companies (claims departments)')

    doc.add_paragraph().add_run('NOT Primary Target:').bold = True
    doc.add_paragraph('• Car owners (well-served, vitamins only)')
    doc.add_paragraph('• VinFast/OEM (won\'t co-develop)')

    add_heading(doc, '9.2 Entry Points (Prioritized)', 2)

    entry_priority = [
        ['Priority', 'Entry Point', 'Rationale', 'Effort'],
        ['1', 'AI Call Center', 'Explicit demand, quick win for leadership', 'Medium'],
        ['2', 'Garage Portal', 'High pain, low politics, clear ROI', 'Medium'],
        ['3', 'Claims Triage', 'Meets 40% AI KPI requirement', 'High'],
    ]
    add_table(doc, entry_priority[0], entry_priority[1:])

    add_heading(doc, '9.3 What NOT to Lead With', 2)

    not_lead = [
        ['Don\'t Lead With', 'Reason'],
        ['Fraud detection', 'Not a leadership priority ("luật ngầm")'],
        ['Consumer claims experience', 'Vitamins, not painkillers'],
        ['VinFast integration', 'They won\'t co-develop'],
        ['Full platform from day one', 'Too much risk, validate first'],
    ]
    add_table(doc, not_lead[0], not_lead[1:])

    add_heading(doc, '9.4 Go-to-Market Strategy', 2)

    doc.add_paragraph().add_run('1. Target insurers, not OEMs').bold = True
    doc.add_paragraph('Insurers have the pain and the budget. OEMs will follow once insurers prove value.')

    doc.add_paragraph().add_run('2. Pitch multiple insurers simultaneously').bold = True
    doc.add_paragraph('Create competitive pressure in the industry. "Your competitor is already piloting."')

    doc.add_paragraph().add_run('3. Focus messaging on:').bold = True
    doc.add_paragraph('• Innovation & modernization (looks good to shareholders)')
    doc.add_paragraph('• Operational efficiency (meets KPIs, saves resources)')
    doc.add_paragraph('• Fast, measurable results (quick wins)')
    doc.add_paragraph('• Personal wins for leadership (recognition, bonuses)')

    doc.add_paragraph().add_run('4. Start with POC, not platform').bold = True
    doc.add_paragraph('5 garages + 1 strong insurer. Prove value before scaling.')

    add_heading(doc, '9.5 Positioning: "Honest Score"', 2)

    doc.add_paragraph(
        'Instead of positioning around fraud detection (negative framing), position around '
        '"Honest Score"—rewarding transparent customers with faster processing. Same AI '
        'technology, better customer psychology, builds trust instead of suspicion.'
    )

    doc.add_page_break()

    # =========================================================================
    # 10. IMPLEMENTATION ROADMAP
    # =========================================================================
    add_heading(doc, '10. Implementation Roadmap', 1)

    add_heading(doc, 'Phase 1: Pilot (2025)', 2)

    phase1 = [
        ['Week', 'Activity', 'Deliverable'],
        ['1-2', 'Identify 1 insurer partner willing to pilot', 'Signed LOI'],
        ['1-2', 'Map 5 target garages (VF authorized + private)', 'Garage list'],
        ['3-4', 'Define MVP scope: Garage Portal OR AI Call Center', 'PRD'],
        ['5-8', 'Build narrow MVP for chosen entry point', 'Working prototype'],
        ['9-12', 'Run POC with garage + insurer pair', 'Pilot data'],
        ['12', 'Evaluate: processing time, satisfaction, ROI', 'Decision to proceed'],
    ]
    add_table(doc, phase1[0], phase1[1:])

    doc.add_paragraph().add_run('Success Criteria:').bold = True
    doc.add_paragraph('• Claims processed 10x faster (target: 4 hours for simple, 48 hours for complex)')
    doc.add_paragraph('• Positive feedback from pilot garages and insurer')
    doc.add_paragraph('• Clear path to revenue model')

    add_heading(doc, 'Phase 2: Expand (2026)', 2)

    phase2 = [
        ['Quarter', 'Activity', 'Target'],
        ['Q1', 'Expand to all VinFast garages (151)', 'Full VF network'],
        ['Q2', 'Add 2 more insurer partners', '3 insurers total'],
        ['Q3', 'Build centralized data center', 'Shared claims database'],
        ['Q4', 'Add second entry point (triage or call center)', 'Multi-product'],
    ]
    add_table(doc, phase2[0], phase2[1:])

    add_heading(doc, 'Phase 3: Scale (2027)', 2)

    phase3 = [
        ['Milestone', 'Description'],
        ['Nationwide coverage', 'Expand beyond VinFast to all auto garages'],
        ['Motorcycle segment', 'Include motorcycle insurance (larger volume)'],
        ['National database', 'Establish industry-standard claims database'],
        ['Platform status', 'Become essential infrastructure for Vietnam collision market'],
    ]
    add_table(doc, phase3[0], phase3[1:])

    add_heading(doc, 'Key Principle', 2)

    doc.add_paragraph().add_run(
        '"Solve a smaller problem first to validate desirability risk. '
        'Don\'t build the platform. Build one feature. Prove it works. Then expand."'
    ).italic = True

    doc.add_page_break()

    # =========================================================================
    # 11. APPENDICES
    # =========================================================================
    add_heading(doc, '11. Appendices', 1)

    add_heading(doc, 'A. Research Documents Index', 2)

    doc_index = [
        ['#', 'File', 'Date', 'Content'],
        ['00', 'timeline-index.md', '-', 'Master synthesis document'],
        ['01', 'problem-statement.md', 'Jul 31', 'BeeVN gaps, opportunities identified'],
        ['02', 'value-proposition.md', 'Aug 2', 'Embedded insurance VP + landing page'],
        ['03', 'embedded-insurance-research.md', 'Aug 6', 'Market context research'],
        ['04', 'market-research.md', 'Aug 7', 'VinFast sales, revenue sizing'],
        ['05', 'pivot-to-claims.md', 'Aug 22-24', 'Pivot decision, competitive landscape'],
        ['06', 'claims-market-research.md', 'Aug 26', 'Market size, fraud types'],
        ['07', 'cto-pitch-mvp.md', 'Aug 29', 'MVP roadmap for technical partner'],
        ['08', 'federated-learning-research.md', 'Sep 22', 'Privacy-preserving ML'],
        ['09', 'claim-triage-system.md', 'Sep 25', 'Triage concept design'],
        ['10', 'customer-interviews.md', 'Sep 24-25', 'Critical interview evidence'],
        ['11', 'garage-interview.md', 'Sep 25', 'Workshop manager insights'],
        ['12', 'marketing-copy.md', 'Oct 10', 'Positioning by audience'],
        ['13', 'product-concept-sep5.md', 'Sep 5', 'Historical product concept'],
        ['14', 'pivot-vpc-bmc.md', 'Aug 22', 'Full VPC/BMC frameworks'],
        ['15', 'weekly-validation-sep30.md', 'Sep 30', 'Synthesis + start small'],
        ['16', 'formal-proposal.md', 'Sep 15', 'Đề án for stakeholders'],
        ['17', 'competitive-analysis.md', 'Sep 18', '12 solutions analyzed'],
    ]
    add_table(doc, doc_index[0], doc_index[1:])

    add_heading(doc, 'B. VinFast Partnership Context', 2)

    doc.add_paragraph(
        'In December 2024, VinFast signed agreements with 7 major insurance companies to '
        'shorten inspection and approval times for VinFast car owners:'
    )

    partners = [
        ['Insurer', 'Type'],
        ['PVI', 'State-owned, largest P&C'],
        ['Bảo Việt', 'State-owned, diversified'],
        ['BIC', 'BIDV subsidiary'],
        ['VBI', 'Vietcombank subsidiary'],
        ['PTI', 'Post & Telecom'],
        ['BSH', 'Private'],
        ['VNI', 'Private'],
    ]
    add_table(doc, partners[0], partners[1:])

    add_heading(doc, 'C. Regulatory Context', 2)

    doc.add_paragraph().add_run('Nghị định 67/2023/NĐ-CP: ').bold = True
    doc.paragraphs[-1].add_run(
        'Regulates documents and responsibilities for mandatory motor vehicle liability insurance. '
        'Directly affects claims documentation checklist requirements.'
    )

    doc.add_paragraph().add_run('Insurance Commission Caps (MoF regulated):').bold = True
    doc.add_paragraph('• Motor liability insurance: max 20%')
    doc.add_paragraph('• Non-life general: 5-20%')
    doc.add_paragraph('• Brokerage: max 15% of premium')

    add_heading(doc, 'D. Glossary', 2)

    glossary = [
        ['Term', 'Definition'],
        ['BeeVN', 'Vingroup\'s in-house insurance broker'],
        ['CQLTT', 'Market regulator (Cục Quản lý Thị trường)'],
        ['DMS', 'Dealer Management System (OEM-provided)'],
        ['MISA', 'Popular Vietnamese accounting software'],
        ['Luật ngầm', '"Unwritten rule" - tacit industry acceptance of fraud'],
        ['VPC', 'Value Proposition Canvas'],
        ['BMC', 'Business Model Canvas'],
        ['POC', 'Proof of Concept'],
        ['TPA', 'Third Party Administrator'],
        ['UBI', 'Usage-Based Insurance'],
    ]
    add_table(doc, glossary[0], glossary[1:])

    # =========================================================================
    # SAVE DOCUMENT
    # =========================================================================
    output_path = '/Users/nntrvy/osr-company/electrik-project/reports/electrik-market-research-report.docx'
    doc.save(output_path)
    print(f'Report saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()
