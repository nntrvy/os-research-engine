#!/usr/bin/env python3
"""
Generate Professional Social Listening Report for Mdm. Pham Lan Anh
Life Insurance Vietnam - Combined VozForums & Facebook Analysis
"""

import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Output directory
OUTPUT_DIR = "/Users/nntrvy/osr-company/os-research-engine/research-outputs/social-listening"

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_chart_complaints_by_company():
    """Bar chart: Complaints by Insurance Company"""
    companies = ['Manulife', 'Prudential', 'AIA', 'FWD', 'Sunlife', 'Generali', 'Chubb', 'Cathay']
    complaints = [85, 70, 65, 45, 40, 30, 20, 15]
    colors = ['#e74c3c', '#e67e22', '#f39c12', '#3498db', '#9b59b6', '#1abc9c', '#34495e', '#95a5a6']

    fig, ax = plt.subplots(figsize=(8, 4))
    bars = ax.barh(companies, complaints, color=colors)
    ax.set_xlabel('Complaint Volume (Relative Score)', fontsize=10)
    ax.set_title('Insurance Companies by Complaint Volume', fontsize=12, fontweight='bold')
    ax.invert_yaxis()

    for bar, val in zip(bars, complaints):
        ax.text(val + 1, bar.get_y() + bar.get_height()/2, str(val), va='center', fontsize=9)

    plt.tight_layout()
    path = f"{OUTPUT_DIR}/chart_complaints.png"
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path

def create_chart_pain_distribution():
    """Pie chart: Pain Point Distribution"""
    labels = ['Misleading Sales\n(Lừa đảo)', 'Claim Denials', 'Hidden Fees', 'Complex Contracts', 'Refund Delays', 'No Accountability']
    sizes = [30, 25, 18, 12, 10, 5]
    colors = ['#e74c3c', '#e67e22', '#f1c40f', '#3498db', '#9b59b6', '#95a5a6']
    explode = (0.05, 0, 0, 0, 0, 0)

    fig, ax = plt.subplots(figsize=(7, 5))
    wedges, texts, autotexts = ax.pie(sizes, explode=explode, labels=labels, colors=colors,
                                       autopct='%1.0f%%', startangle=90,
                                       textprops={'fontsize': 9})
    ax.set_title('Customer Pain Points Distribution', fontsize=12, fontweight='bold')

    plt.tight_layout()
    path = f"{OUTPUT_DIR}/chart_pains.png"
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path

def create_chart_market_stats():
    """Bar chart: Key Market Statistics"""
    metrics = ['Contracts Voided\n(2023)', 'New Policy\nDecline (2024)', 'First-Year\nCancellation', 'Principal Lost on\nEarly Exit']
    values = [3.0, 19, 25, 30]
    units = ['million', '%', '%', '%']
    colors = ['#e74c3c', '#e67e22', '#f39c12', '#c0392b']

    fig, ax = plt.subplots(figsize=(8, 4))
    bars = ax.bar(metrics, values, color=colors, width=0.6)
    ax.set_ylabel('Value', fontsize=10)
    ax.set_title('Vietnam Life Insurance Market Crisis (2023-2024)', fontsize=12, fontweight='bold')

    for bar, val, unit in zip(bars, values, units):
        label = f"{val} {unit}" if unit != 'million' else f"{val}M"
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, label,
                ha='center', va='bottom', fontsize=10, fontweight='bold')

    ax.set_ylim(0, max(values) * 1.2)
    plt.tight_layout()
    path = f"{OUTPUT_DIR}/chart_stats.png"
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path

def create_chart_sentiment():
    """Horizontal bar: Sentiment Keywords"""
    keywords = ['Lừa đảo (Fraud)', 'Im lặng (Silence)', 'Mập mờ (Vague)', 'Mệt mỏi (Exhausted)', 'Chó má (Cursing)']
    mentions = [50, 30, 25, 20, 15]
    colors = ['#c0392b', '#7f8c8d', '#8e44ad', '#d35400', '#2c3e50']

    fig, ax = plt.subplots(figsize=(7, 3.5))
    bars = ax.barh(keywords, mentions, color=colors)
    ax.set_xlabel('Mentions', fontsize=10)
    ax.set_title('Negative Sentiment Keywords (Facebook)', fontsize=12, fontweight='bold')
    ax.invert_yaxis()

    for bar, val in zip(bars, mentions):
        ax.text(val + 0.5, bar.get_y() + bar.get_height()/2, f'{val}+', va='center', fontsize=9)

    plt.tight_layout()
    path = f"{OUTPUT_DIR}/chart_sentiment.png"
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path

def create_document():
    """Create the Word document"""
    doc = Document()

    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ===== TITLE PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title_run = title.add_run("SOCIAL LISTENING REPORT")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(44, 62, 80)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Life Insurance Market in Vietnam")
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(127, 140, 141)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Info box
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run("Prepared for: Mdm. Pham Lan Anh\nDate: 04.12.25\nAnalyst: OS Research Engine")
    info_run.font.size = Pt(12)
    info_run.font.color.rgb = RGBColor(52, 73, 94)

    doc.add_paragraph()
    doc.add_paragraph()

    sources = doc.add_paragraph()
    sources.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src_run = sources.add_run("Sources: VozForums • Facebook Groups • Vietnamese News")
    src_run.font.size = Pt(11)
    src_run.font.italic = True
    src_run.font.color.rgb = RGBColor(149, 165, 166)

    doc.add_page_break()

    # ===== EXECUTIVE SUMMARY =====
    h1 = doc.add_heading('Executive Summary', level=1)
    h1.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    summary_text = """Vietnamese consumers deeply distrust life insurance due to widespread deceptive sales practices. This report analyzes 35+ sources across VozForums and Facebook, revealing systemic issues that have led to 3+ million contract cancellations in 2023 alone.

Key Finding: The dominant complaint is misleading sales tactics, with customers losing 15-50% of their money when canceling policies they didn't fully understand."""

    p = doc.add_paragraph(summary_text)
    p.paragraph_format.space_after = Pt(12)

    # Key Numbers Table
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['3M+', '19%', '100 pages', '74 years']
    labels = ['Contracts Voided', 'Market Decline', 'Avg. Contract', 'Max Term Found']

    for i, (header, label) in enumerate(zip(headers, labels)):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(16)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(231, 76, 60)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'FDF2F2')

        cell2 = table.rows[1].cells[i]
        cell2.text = label
        cell2.paragraphs[0].runs[0].font.size = Pt(9)
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell2, 'FDF2F2')

    doc.add_paragraph()

    # ===== MARKET CRISIS CHART =====
    h2 = doc.add_heading('Market Crisis at a Glance', level=1)
    h2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    chart_path = create_chart_market_stats()
    doc.add_picture(chart_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ===== TOP CUSTOMER PAINS =====
    h3 = doc.add_heading('Top Customer Pains', level=1)
    h3.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    chart_path = create_chart_pain_distribution()
    doc.add_picture(chart_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Pain details table
    pain_table = doc.add_table(rows=5, cols=3)
    pain_table.style = 'Table Grid'

    pain_data = [
        ('Misleading Sales', 'HIGH', '50+ mentions - agents lie about terms'),
        ('Claim Denials', 'HIGH', 'Claims rejected despite coverage'),
        ('Hidden Fees', 'HIGH', 'Annual deductions not explained'),
        ('Complex Contracts', 'MEDIUM', '100 pages nobody understands'),
        ('Refund Delays', 'MEDIUM', '3+ months wait for legal cancellations'),
    ]

    for i, (pain, severity, detail) in enumerate(pain_data):
        pain_table.rows[i].cells[0].text = pain
        pain_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        sev_cell = pain_table.rows[i].cells[1]
        sev_cell.text = severity
        sev_cell.paragraphs[0].runs[0].bold = True
        if severity == 'HIGH':
            sev_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(231, 76, 60)

        pain_table.rows[i].cells[2].text = detail

    doc.add_page_break()

    # ===== REAL CUSTOMER VOICES =====
    h4 = doc.add_heading('Real Customer Voices', level=1)
    h4.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    intro = doc.add_paragraph("Direct quotes from Vietnamese consumers (translated):")
    intro.runs[0].font.italic = True

    quotes = [
        ('"The contract is nearly 100 pages, even finance professors and PhDs may not fully understand it"', 'VozForums - Contract Complexity'),
        ('"FWD insurance fraud. Customers beware. They claim the electronic contract has exclusions that the paper contract doesn\'t"', 'Facebook - FWD Complaint'),
        ('"Waiting for Manulife is exhausting... when healthy I bought insurance, now sick and facing risks from this damn company"', 'Facebook - Manulife Brain Tumor Case'),
        ('"I\'ve been demanding my refund for over 3 months now. Still silence."', 'Facebook - Sunlife Cancellation'),
        ('"Anyone who becomes insurance sales, I automatically unfriend"', 'VozForums - Social Impact'),
    ]

    for quote, source in quotes:
        q = doc.add_paragraph()
        q_run = q.add_run(quote)
        q_run.font.size = Pt(11)
        q_run.italic = True
        q.paragraph_format.left_indent = Inches(0.3)
        q.paragraph_format.space_after = Pt(2)

        s = doc.add_paragraph()
        s_run = s.add_run(f"— {source}")
        s_run.font.size = Pt(9)
        s_run.font.color.rgb = RGBColor(127, 140, 141)
        s.paragraph_format.left_indent = Inches(0.3)
        s.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # ===== COMPANIES ANALYSIS =====
    h5 = doc.add_heading('Companies by Complaint Volume', level=1)
    h5.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    chart_path = create_chart_complaints_by_company()
    doc.add_picture(chart_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Sentiment chart
    h6 = doc.add_heading('Sentiment Analysis (Facebook)', level=2)
    h6.runs[0].font.color.rgb = RGBColor(52, 73, 94)

    chart_path = create_chart_sentiment()
    doc.add_picture(chart_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ===== WORKAROUNDS =====
    h7 = doc.add_heading('How Customers Are Coping', level=1)
    h7.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    workarounds = [
        ('🏦 Government Insurance Only', 'Avoiding all private insurers, using only BHYT/BHXH'),
        ('🥇 Buying Gold Instead', 'Physical gold as alternative to insurance savings'),
        ('👥 Joining Support Groups', 'Facebook groups with 50K-123K members for advice'),
        ('🚫 Cutting Social Ties', 'Unfriending anyone who sells insurance'),
        ('💸 Accepting Losses', 'Canceling early despite losing 15-50% of premiums'),
    ]

    work_table = doc.add_table(rows=5, cols=2)
    work_table.style = 'Table Grid'

    for i, (title, desc) in enumerate(workarounds):
        work_table.rows[i].cells[0].text = title
        work_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        work_table.rows[i].cells[0].width = Inches(2)
        work_table.rows[i].cells[1].text = desc

    doc.add_paragraph()

    # ===== OPPORTUNITIES =====
    h8 = doc.add_heading('Market Opportunities', level=1)
    h8.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    opps = [
        ('Transparent Insurance Products', 'Simple 1-2 page contracts with clear terms, no hidden fees'),
        ('Independent Advisory Platform', 'Neutral reviews and contract analysis (not agent-driven)'),
        ('Digital-First Distribution', 'Online platform bypassing traditional agent networks'),
        ('Dispute Resolution Service', 'Mediation for claim denials and refund issues'),
    ]

    for title, desc in opps:
        p = doc.add_paragraph()
        p_run = p.add_run(f"✓ {title}: ")
        p_run.bold = True
        p_run.font.color.rgb = RGBColor(39, 174, 96)
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

    # ===== NEXT STEPS =====
    h9 = doc.add_heading('Recommended Next Steps', level=1)
    h9.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    steps = [
        'Interview 10 people who canceled insurance in 2023-2024',
        'Landing page test: "Would you buy insurance with transparent terms?"',
        'Wizard of Oz: Manual contract review service pilot',
        'Survey: Willingness to pay for independent insurance advisor',
    ]

    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p_run = p.add_run(f"{i}. {step}")
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer.add_run("— End of Report —\nOS Research Engine • December 2025")
    f_run.font.size = Pt(10)
    f_run.font.color.rgb = RGBColor(149, 165, 166)
    f_run.italic = True

    # Save
    output_path = f"{OUTPUT_DIR}/04.12.25 - Life Insurance Vietnam Report - Mdm Pham Lan Anh.docx"
    doc.save(output_path)
    print(f"Report saved: {output_path}")
    return output_path

if __name__ == "__main__":
    create_document()
