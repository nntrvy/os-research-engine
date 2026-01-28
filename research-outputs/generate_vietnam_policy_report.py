#!/usr/bin/env python3
"""
Generate Vietnam Innovation & Entrepreneurship Policy Summary 2025
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_footer(doc, footer_text):
    """Add footer to all pages"""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = footer_text
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

def create_report():
    """Create the Vietnam Innovation Policy Report"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Vietnam Innovation & Entrepreneurship Policy Summary 2025', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()  # Spacing

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', 1)

    summary_text = """Vietnam has positioned innovation, science, and technology as core drivers of national development through a comprehensive policy framework launched in late 2024 and early 2025. The government aims to transform Vietnam into a high-income nation by 2045, with digital economy accounting for 50% of GDP. Key initiatives include substantial R&D investment targets (2% of GDP by 2030), creation of specialized investment funds for emerging technologies, and systematic reduction of administrative barriers for entrepreneurs. The ecosystem has shown strong momentum with over 4,000 active startups, 2 unicorns, and USD 2.3 billion in venture capital deployed in 2024 alone."""

    para = doc.add_paragraph(summary_text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 2. Key Resolutions
    doc.add_heading('2. Key Government Resolutions', 1)

    # Resolution 57-NQ/TW
    doc.add_heading('Resolution 57-NQ/TW (December 22, 2024)', 2)
    doc.add_paragraph('Establishes science, technology, and innovation as fundamental drivers of national socio-economic growth', style='List Bullet')
    doc.add_paragraph('Target: 30% of GDP from digital economy by 2030', style='List Bullet')
    doc.add_paragraph('Target: 2% of GDP allocated to R&D investment by 2030', style='List Bullet')
    doc.add_paragraph('Goal: Rank in Top 3 ASEAN countries for innovation by 2030', style='List Bullet')
    doc.add_paragraph('Goal: Rank in Top 40 globally in innovation indices by 2030', style='List Bullet')

    # Resolution 59-NQ/TW
    doc.add_heading('Resolution 59-NQ/TW (January 24, 2025)', 2)
    doc.add_paragraph('Focuses on global integration and intellectual property commercialization', style='List Bullet')
    doc.add_paragraph('Target: 30% reduction in administrative burden for businesses by 2025', style='List Bullet')
    doc.add_paragraph('Establishes framework for international R&D collaboration', style='List Bullet')
    doc.add_paragraph('Promotes technology transfer from research institutions to commercial entities', style='List Bullet')

    # 3. Startup Ecosystem Statistics
    doc.add_heading('3. Startup Ecosystem Statistics (2024)', 1)

    doc.add_paragraph('Total active startups: 4,000+', style='List Bullet')
    doc.add_paragraph('Unicorns (USD 1B+ valuation): 2 companies', style='List Bullet')
    doc.add_paragraph('High-value companies (USD 100M+ valuation): 11 companies', style='List Bullet')
    doc.add_paragraph('Total venture capital invested: USD 2.3 billion across 141 deals', style='List Bullet')
    doc.add_paragraph('AI/ML startups: 765 companies (25% of total tech startup ecosystem)', style='List Bullet')
    doc.add_paragraph('Startup support organizations: 1,400+', style='List Bullet')
    doc.add_paragraph('Co-working spaces: 200+', style='List Bullet')
    doc.add_paragraph('Venture capital funds: 208', style='List Bullet')
    doc.add_paragraph('Business accelerators: 35', style='List Bullet')
    doc.add_paragraph('Incubators: 80', style='List Bullet')

    # 4. Key Legislation & Financial Incentives
    doc.add_heading('4. Key Legislation & Financial Incentives', 1)

    # Project 844
    doc.add_heading('Project 844: National Startup Support Program', 2)
    doc.add_paragraph('Funding allocation: USD 42 million', style='List Bullet')
    doc.add_paragraph('Targeted support for 800 startup projects through 2025', style='List Bullet')
    doc.add_paragraph('Covers seed funding, mentorship, and market access programs', style='List Bullet')

    # Decree 182/2024
    doc.add_heading('Decree 182/2024: Investment Support Fund', 2)
    doc.add_paragraph('Dedicated funding for semiconductor research and development', style='List Bullet')
    doc.add_paragraph('Dedicated funding for artificial intelligence research and development', style='List Bullet')
    doc.add_paragraph('Supports domestic chip design and fabrication capabilities', style='List Bullet')

    # Resolution 193/2025
    doc.add_heading('Resolution 193/2025: Risk Acceptance Framework', 2)
    doc.add_paragraph('Protects R&D personnel from liability for experimental failures', style='List Bullet')
    doc.add_paragraph('Encourages bold innovation through legal safeguards', style='List Bullet')
    doc.add_paragraph('Establishes clear criteria for acceptable vs. negligent risk-taking', style='List Bullet')

    # Decree 160/2025
    doc.add_heading('Decree 160/2025: AI Development Fund', 2)
    doc.add_paragraph('Fund size: 1 trillion VND (approximately USD 40 million)', style='List Bullet')
    doc.add_paragraph('Focus areas: AI research, product development, and commercialization', style='List Bullet')
    doc.add_paragraph('Available to universities, research institutes, and private companies', style='List Bullet')

    # National Venture Capital Fund
    doc.add_heading('National Venture Capital Fund (2025)', 2)
    doc.add_paragraph('State-backed venture capital fund launched in 2025', style='List Bullet')
    doc.add_paragraph('Co-investment model with private sector venture funds', style='List Bullet')
    doc.add_paragraph('Targets early-stage and growth-stage technology companies', style='List Bullet')

    # Tax Incentives
    doc.add_heading('Tax Incentives for Innovation Enterprises', 2)
    doc.add_paragraph('Corporate Income Tax (CIT): Preferential rate of 10% for 30 years', style='List Bullet')
    doc.add_paragraph('Tax holiday: 4-year complete exemption from CIT', style='List Bullet')
    doc.add_paragraph('Tax reduction: 50% CIT reduction for subsequent 9 years', style='List Bullet')
    doc.add_paragraph('Applies to qualified technology enterprises and innovation centers', style='List Bullet')

    # 5. Talent Attraction Initiatives
    doc.add_heading('5. Talent Attraction & Immigration Reform', 1)

    # 5-Year Talent Visa
    doc.add_heading('5-Year Talent Visa Program (August 2025)', 2)
    doc.add_paragraph('Eligibility: Top-tier academics, researchers, and digital technology specialists', style='List Bullet')
    doc.add_paragraph('Duration: 5-year renewable visa with streamlined application process', style='List Bullet')
    doc.add_paragraph('Benefits: Family sponsorship, property ownership rights, work flexibility', style='List Bullet')

    # Work Permit Reforms
    doc.add_heading('Work Permit Exemptions & Streamlining', 2)
    doc.add_paragraph('Work permit exemptions for foreign experts at National Innovation Center', style='List Bullet')
    doc.add_paragraph('Experience requirements reduced by 33-40% for technology roles', style='List Bullet')
    doc.add_paragraph('Fast-track processing for STEM professionals and startup founders', style='List Bullet')

    # 6. Open Source Policy
    doc.add_heading('6. Open Source & Domestic Technology Promotion', 1)

    # Circular 34/2025
    doc.add_heading('Circular 34/2025: Government Procurement Preferences', 2)
    doc.add_paragraph('Preferential treatment for Vietnamese open-source software products', style='List Bullet')
    doc.add_paragraph('Requirement for government agencies to consider domestic solutions first', style='List Bullet')
    doc.add_paragraph('Supports local software development and reduces foreign technology dependence', style='List Bullet')
    doc.add_paragraph('Promotes transparency and security through open-source adoption', style='List Bullet')

    # 7. Timeline & Strategic Targets
    doc.add_heading('7. Strategic Timeline & National Targets', 1)

    # 2025 Targets
    doc.add_heading('2025 Milestones', 2)
    doc.add_paragraph('Digital economy: 20% of GDP', style='List Bullet')
    doc.add_paragraph('Launch National Venture Capital Fund', style='List Bullet')
    doc.add_paragraph('30% reduction in administrative burden for businesses', style='List Bullet')
    doc.add_paragraph('5-Year Talent Visa program operational', style='List Bullet')

    # 2027 Targets
    doc.add_heading('2027 Milestones', 2)
    doc.add_paragraph('Semiconductor self-reliance in chip design and production', style='List Bullet')
    doc.add_paragraph('Operational domestic semiconductor fabrication facilities', style='List Bullet')

    # 2030 Targets
    doc.add_heading('2030 Vision', 2)
    doc.add_paragraph('Digital economy: 30% of GDP', style='List Bullet')
    doc.add_paragraph('R&D investment: 2% of GDP', style='List Bullet')
    doc.add_paragraph('Innovation ranking: Top 3 in ASEAN', style='List Bullet')
    doc.add_paragraph('Innovation ranking: Top 40 globally', style='List Bullet')

    # 2045 Long-term Vision
    doc.add_heading('2045 Long-term Vision', 2)
    doc.add_paragraph('Digital economy: 50% of GDP', style='List Bullet')
    doc.add_paragraph('National status: High-income developed nation', style='List Bullet')
    doc.add_paragraph('Global innovation leadership in Southeast Asia', style='List Bullet')

    # Add footer
    add_footer(doc, "Source: OS Research Engine | January 2025")

    # Save document
    output_path = '/Users/Vincent/OSR-company/os-research-engine/research-outputs/vietnam-innovation-policy-summary-2025.docx'
    doc.save(output_path)
    print(f"Report successfully created: {output_path}")

if __name__ == '__main__':
    create_report()
