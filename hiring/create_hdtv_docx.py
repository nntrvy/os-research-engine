"""
Generate Hợp đồng thử việc - Content Strategist (.docx)
Vietnamese contract format with proper styling.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3

# ── Helper functions ──

def add_blank(count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

def add_centered(text, bold=False, size=13, color=None, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_heading_styled(text, size=13, bold=True, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_body(text, bold=False, indent_left=0, space_after=4, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent_left:
        p.paragraph_format.left_indent = Cm(indent_left)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = bold
    run.italic = italic
    return p

def add_mixed(parts, indent_left=0, space_after=4):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent_left:
        p.paragraph_format.left_indent = Cm(indent_left)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.bold = bold
        run.italic = italic
    return p

def add_field(label, value=".......................................", indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run_label = p.add_run(label)
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(13)
    run_label.bold = False
    run_val = p.add_run(value)
    run_val.font.name = 'Times New Roman'
    run_val.font.size = Pt(13)
    return p

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold

def add_simple_table(rows_data, col_widths=None, header=True):
    """rows_data = list of lists. First row is header if header=True."""
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Style borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            is_header = (i == 0 and header)
            set_cell_text(cell, cell_text, bold=is_header, size=12)
            if is_header:
                set_cell_shading(cell, "D9E2F3")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table


# ═══════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ═══════════════════════════════════════════════════════════════

# ── National header ──
add_centered("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, size=13, space_after=0)
add_centered("Độc lập – Tự do – Hạnh phúc", bold=True, size=13, space_after=0)
add_centered("─────────  ✦  ─────────", size=12, space_after=12)

# ── Company header (left-aligned) ──
add_body("CÔNG TY ............................................................", bold=True, space_after=2)
add_body("Số: ......../HĐTV/2026", space_after=12)

# ── Title ──
add_centered("HỢP ĐỒNG THỬ VIỆC", bold=True, size=16, space_after=4)
add_centered("(Vị trí: Content Strategist – Chuyên viên Chiến lược Nội dung)", size=13, space_after=12)

# ── Legal basis ──
add_body("Căn cứ pháp lý:", bold=True, space_after=2)
add_body("–  Bộ luật Lao động số 45/2019/QH14 ngày 20 tháng 11 năm 2019;", indent_left=0.5, space_after=2)
add_body("–  Nghị định số 145/2020/NĐ-CP ngày 14/12/2020 của Chính phủ quy định chi tiết và hướng dẫn thi hành một số điều của Bộ luật Lao động;", indent_left=0.5, space_after=2)
add_body("–  Luật Bảo hiểm xã hội số 41/2024/QH15 ngày 29 tháng 6 năm 2024;", indent_left=0.5, space_after=2)
add_body("–  Thỏa thuận của hai bên.", indent_left=0.5, space_after=8)

add_body("Hôm nay, ngày ...... tháng ...... năm 2026, tại .................................................................", space_after=6)
add_body("Chúng tôi gồm:", bold=True, space_after=8)

# ══════════ BÊN A ══════════
add_heading_styled("BÊN A: NGƯỜI SỬ DỤNG LAO ĐỘNG", size=13)
add_field("Tên đơn vị:  ")
add_field("Địa chỉ trụ sở:  ")
add_field("Mã số thuế:  ")
add_field("Điện thoại:  ")
add_field("Đại diện bởi:  ", value="Ông/Bà .............................................................")
add_field("Chức vụ:  ")

add_body("(Sau đây gọi tắt là \"Bên A\")", italic=True, space_after=10)

# ══════════ BÊN B ══════════
add_heading_styled("BÊN B: NGƯỜI LAO ĐỘNG", size=13)
add_field("Họ và tên:  ")
add_field("Ngày sinh:  ", value="......./......./............          Giới tính: ................")
add_field("Số CCCD/CMND:  ")
add_field("Ngày cấp:  ", value="......./......./............          Nơi cấp: ................................")
add_field("Địa chỉ thường trú:  ")
add_field("Địa chỉ hiện tại:  ")
add_field("Số điện thoại:  ")
add_field("Email:  ")
add_field("Trình độ học vấn:  ")

add_body("(Sau đây gọi tắt là \"Bên B\")", italic=True, space_after=10)

add_body("Hai bên cùng thống nhất ký kết Hợp đồng thử việc (sau đây gọi là \"Hợp đồng\") với các điều khoản sau:", space_after=10)

# ══════════════════════════════════════════════════════════════
# ĐIỀU 1
# ══════════════════════════════════════════════════════════════
add_heading_styled("ĐIỀU 1: CÔNG VIỆC, ĐỊA ĐIỂM VÀ THỜI GIAN THỬ VIỆC", size=13)

add_mixed([
    ("1.1. Vị trí công việc: ", True, False),
    ("Content Strategist (Chuyên viên Chiến lược Nội dung)", False, False),
], indent_left=0.5, space_after=3)

add_mixed([
    ("1.2. Phòng ban: ", True, False),
    (".......................................", False, False),
], indent_left=0.5, space_after=3)

add_mixed([
    ("1.3. Địa điểm làm việc: ", True, False),
    (".......................................", False, False),
], indent_left=0.5, space_after=3)

add_mixed([
    ("1.4. Loại hợp đồng: ", True, False),
    ("Hợp đồng thử việc", False, False),
], indent_left=0.5, space_after=3)

add_mixed([
    ("1.5. Thời gian thử việc: ", True, False),
    ("02 (hai) tháng", False, False),
], indent_left=0.5, space_after=2)
add_body("Từ ngày: ......./......./2026    đến ngày: ......./......./2026", indent_left=1, space_after=2)
add_body("(Theo Điều 25 khoản 2 BLLĐ 2019: Thời gian thử việc không quá 60 ngày đối với công việc có chức danh nghề cần trình độ chuyên môn, kỹ thuật từ cao đẳng trở lên.)", indent_left=1, space_after=6, italic=True)

add_mixed([
    ("1.6. Thời giờ làm việc: ", True, False),
    ("......... giờ/tuần", False, False),
], indent_left=0.5, space_after=2)
add_body("Từ Thứ Hai đến Thứ Sáu, từ ......... giờ đến ......... giờ", indent_left=1, space_after=2)
add_body("Nghỉ trưa: từ ......... giờ đến ......... giờ", indent_left=1, space_after=2)
add_body("Hình thức:  ☐ Toàn thời gian tại văn phòng    ☐ Hybrid    ☐ Từ xa (Remote)", indent_left=1, space_after=8)

# ══════════════════════════════════════════════════════════════
# ĐIỀU 2
# ══════════════════════════════════════════════════════════════
add_heading_styled("ĐIỀU 2: MÔ TẢ CÔNG VIỆC VÀ TIÊU CHÍ ĐÁNH GIÁ", size=13)

add_body("2.1. Công việc chính:", bold=True, indent_left=0.5, space_after=3)

tasks = [
    "Xây dựng và triển khai chiến lược nội dung (content strategy) tổng thể cho các kênh truyền thông của Công ty;",
    "Nghiên cứu đối tượng mục tiêu, phân tích thị trường và xu hướng nội dung để đề xuất hướng đi phù hợp;",
    "Lập kế hoạch nội dung (content calendar) theo tuần, tháng, quý;",
    "Sáng tạo, biên tập và quản lý nội dung trên các nền tảng: website, mạng xã hội, email marketing và các kênh khác theo yêu cầu;",
    "Đo lường hiệu quả nội dung thông qua các chỉ số KPI đã thống nhất; đề xuất tối ưu dựa trên dữ liệu;",
    "Phối hợp với các bộ phận liên quan (Design, Product, Marketing) để đảm bảo tính nhất quán của thông điệp thương hiệu;",
    "Các công việc khác theo sự phân công hợp lý của quản lý trực tiếp, phù hợp với vị trí.",
]
for i, task in enumerate(tasks):
    label = chr(97 + i)  # a, b, c...
    add_body(f"{label})  {task}", indent_left=1, space_after=2)

add_blank()
add_body("2.2. Tiêu chí đánh giá thử việc:", bold=True, indent_left=0.5, space_after=4)

add_body("Tiêu chí đánh giá cốt lõi trong suốt thời gian thử việc là khả năng ghi chép, tổng hợp và báo cáo một cách có hệ thống. Cụ thể, Bên B phải đảm bảo tất cả các hoạt động công việc được lưu trữ đầy đủ, rõ ràng và có thể truy xuất được.", indent_left=0.5, space_after=6)

eval_data = [
    ["STT", "Tiêu chí đánh giá", "Yêu cầu cụ thể"],
    ["1",
     "Ghi nhận câu hỏi (Questions Asked)\n"
     "Mọi câu hỏi nghiên cứu, câu hỏi đặt ra cho khách hàng, đối tác, nội bộ đều được ghi chép lại.",
     "100% câu hỏi được lưu trữ có hệ thống, gắn ngày tháng và ngữ cảnh."],
    ["2",
     "Thu thập bằng chứng (Evidence Collected)\n"
     "Mọi dữ liệu, số liệu, phản hồi, screenshot, link tham khảo, kết quả khảo sát đều được lưu trữ.",
     "100% bằng chứng được phân loại, gắn nguồn gốc, dễ truy xuất."],
    ["3",
     "Tổng hợp insight (Insights Gained)\n"
     "Từ bằng chứng thu thập, phải rút ra được insight rõ ràng, có căn cứ, và ghi nhận đầy đủ.",
     "Mỗi insight phải gắn liền với bằng chứng cụ thể. Không insight nào thiếu dữ liệu hỗ trợ."],
    ["4",
     "Ghi nhận hành động (Actions Taken)\n"
     "Mọi hành động đã thực hiện, quyết định đã đưa ra, thay đổi đã áp dụng đều được ghi chép.",
     "100% action items được ghi nhận kèm kết quả và trạng thái (đã làm / đang làm / chưa làm)."],
    ["5",
     "Báo cáo Loom hàng tuần (Weekly Loom Report)\n"
     "Mỗi tuần, Bên B phải trình bày tổng hợp công việc bằng video Loom gửi cho quản lý trực tiếp.",
     "Nộp đúng hạn mỗi tuần. Nội dung bao gồm: Câu hỏi → Bằng chứng → Insight → Hành động. Rõ ràng, có cấu trúc."],
]
add_simple_table(eval_data, col_widths=[1, 6.5, 6])

add_blank()
add_body("Lưu ý: Tiêu chí đánh giá không dựa trên khối lượng nội dung sản xuất, mà dựa trên chất lượng ghi chép, mức độ hệ thống hóa và tính nhất quán của báo cáo hàng tuần.", indent_left=0.5, italic=True, space_after=3)

add_body("2.3. Quy định về báo cáo Loom hàng tuần:", bold=True, indent_left=0.5, space_after=4)

loom_data = [
    ["Nội dung", "Chi tiết"],
    ["Hình thức", "Video trình bày qua Loom (có chia sẻ màn hình)"],
    ["Tần suất", "01 lần/tuần, nộp vào ngày ......... hàng tuần"],
    ["Thời lượng", "Tối thiểu 5 phút, tối đa 15 phút"],
    ["Cấu trúc bắt buộc",
     "1. Questions: Các câu hỏi đã đặt ra trong tuần\n"
     "2. Evidence: Bằng chứng/dữ liệu đã thu thập\n"
     "3. Insights: Nhận định rút ra từ bằng chứng\n"
     "4. Actions: Hành động đã thực hiện & kế hoạch tuần tiếp theo"],
    ["Gửi đến", "Quản lý trực tiếp qua email/Slack/công cụ nội bộ"],
    ["Lưu trữ", "Tất cả video Loom được lưu trong thư mục chung do Bên A chỉ định"],
]
add_simple_table(loom_data, col_widths=[3.5, 10])
add_blank()

# ══════════════════════════════════════════════════════════════
# ĐIỀU 3
# ══════════════════════════════════════════════════════════════
add_heading_styled("ĐIỀU 3: TIỀN LƯƠNG VÀ PHƯƠNG THỨC THANH TOÁN", size=13)

add_mixed([
    ("3.1. Mức lương thử việc: ", True, False),
    (".............................. VNĐ/tháng", False, False),
], indent_left=0.5, space_after=2)
add_body("(Bằng chữ: .............................................................................................................)", indent_left=1, space_after=3)
add_body("(Theo Điều 26 BLLĐ 2019: Tiền lương thử việc ít nhất bằng 85% mức lương của công việc đó.)", indent_left=1, italic=True, space_after=5)

add_mixed([
    ("3.2. Mức lương chính thức (dự kiến sau thử việc): ", True, False),
    (".............................. VNĐ/tháng", False, False),
], indent_left=0.5, space_after=5)

add_body("3.3. Phương thức thanh toán:", bold=True, indent_left=0.5, space_after=2)
add_body("Hình thức: Chuyển khoản ngân hàng", indent_left=1, space_after=2)
add_field("Ngân hàng:  ", indent=1)
add_field("Số tài khoản:  ", indent=1)
add_field("Chủ tài khoản:  ", indent=1)
add_field("Kỳ thanh toán:  ", value="Ngày ......... hàng tháng", indent=1)

add_blank()
add_body("3.4. Phụ cấp (nếu có):", bold=True, indent_left=0.5, space_after=4)

allowance_data = [
    ["Loại phụ cấp", "Số tiền (VNĐ/tháng)"],
    ["Phụ cấp ăn trưa", ""],
    ["Phụ cấp đi lại / xăng xe", ""],
    ["Phụ cấp điện thoại", ""],
    ["Khác: .........................", ""],
]
add_simple_table(allowance_data, col_widths=[8, 5])
add_blank()

# ══════════════════════════════════════════════════════════════
# ĐIỀU 4
# ══════════════════════════════════════════════════════════════
add_heading_styled("ĐIỀU 4: QUYỀN VÀ NGHĨA VỤ CỦA BÊN A", size=13)

add_body("4.1. Quyền của Bên A:", bold=True, indent_left=0.5, space_after=3)
rights_a = [
    "Quản lý, điều hành, phân công công việc cho Bên B theo đúng vị trí và chức năng đã thỏa thuận;",
    "Đánh giá năng lực và kết quả làm việc của Bên B trong thời gian thử việc;",
    "Chấm dứt Hợp đồng thử việc theo quy định tại Điều 7 của Hợp đồng này;",
    "Khen thưởng, kỷ luật Bên B theo nội quy lao động của Công ty.",
]
for i, item in enumerate(rights_a):
    add_body(f"{chr(97+i)})  {item}", indent_left=1, space_after=2)

add_blank()
add_body("4.2. Nghĩa vụ của Bên A:", bold=True, indent_left=0.5, space_after=3)
obligations_a = [
    "Bố trí công việc đúng theo nội dung đã thỏa thuận trong Hợp đồng;",
    "Thanh toán đầy đủ, đúng hạn tiền lương và các chế độ cho Bên B;",
    "Cung cấp điều kiện, phương tiện làm việc cần thiết (máy tính, tài khoản phần mềm, tài liệu nội bộ...);",
    "Hướng dẫn, đào tạo ban đầu để Bên B nắm bắt công việc và văn hóa công ty;",
    "Thông báo kết quả đánh giá thử việc cho Bên B theo Điều 27 Bộ luật Lao động 2019.",
]
for i, item in enumerate(obligations_a):
    add_body(f"{chr(97+i)})  {item}", indent_left=1, space_after=2)

# ══════════════════════════════════════════════════════════════
# ĐIỀU 5
# ══════════════════════════════════════════════════════════════
add_heading_styled("ĐIỀU 5: QUYỀN VÀ NGHĨA VỤ CỦA BÊN B", size=13)

add_body("5.1. Quyền của Bên B:", bold=True, indent_left=0.5, space_after=3)
rights_b = [
    "Được cung cấp đầy đủ điều kiện, phương tiện để thực hiện công việc;",
    "Được hưởng lương và các chế độ theo thỏa thuận trong Hợp đồng này;",
    "Được đào tạo, hướng dẫn để hoàn thành nhiệm vụ;",
    "Chấm dứt Hợp đồng thử việc theo quy định tại Điều 7 của Hợp đồng này;",
    "Khiếu nại, tố cáo theo quy định pháp luật khi quyền lợi bị xâm phạm.",
]
for i, item in enumerate(rights_b):
    add_body(f"{chr(97+i)})  {item}", indent_left=1, space_after=2)

add_blank()
add_body("5.2. Nghĩa vụ của Bên B:", bold=True, indent_left=0.5, space_after=3)
obligations_b = [
    "Hoàn thành công việc được giao đúng tiến độ và chất lượng;",
    "Chấp hành nội quy lao động, quy chế và các quy định của Công ty;",
    "Tuân thủ sự quản lý, điều hành và phân công công việc hợp lý của Bên A;",
    "Bảo mật thông tin kinh doanh, bí mật thương mại và dữ liệu nội bộ của Công ty (xem Điều 8);",
    "Bồi thường thiệt hại nếu vi phạm quy định bảo mật hoặc gây thiệt hại do lỗi cố ý.",
]
for i, item in enumerate(obligations_b):
    add_body(f"{chr(97+i)})  {item}", indent_left=1, space_after=2)

# ══════════════════════════════════════════════════════════════
# ĐIỀU 6
# ══════════════════════════════════════════════════════════════
add_heading_styled("ĐIỀU 6: BẢO HIỂM XÃ HỘI VÀ CÁC CHẾ ĐỘ PHÚC LỢI", size=13)

add_body("6.1. Bảo hiểm xã hội, bảo hiểm y tế, bảo hiểm thất nghiệp:", bold=True, indent_left=0.5, space_after=3)
add_body("Trong thời gian thử việc, Bên A và Bên B thực hiện đóng bảo hiểm xã hội, bảo hiểm y tế, bảo hiểm thất nghiệp theo quy định của pháp luật hiện hành.", indent_left=0.5, space_after=3)
add_body("(Theo Luật Bảo hiểm xã hội 2024, người lao động trong thời gian thử việc thuộc đối tượng tham gia BHXH bắt buộc. Bên A có trách nhiệm đóng BHXH cho Bên B kể từ ngày bắt đầu thử việc.)", indent_left=0.5, italic=True, space_after=5)

add_body("6.2. Các chế độ khác (nếu có):", bold=True, indent_left=0.5, space_after=3)
add_body("☐  Bảo hiểm sức khỏe bổ sung: ...................................................................................", indent_left=1, space_after=2)
add_body("☐  Nghỉ phép trong thời gian thử việc: Theo thỏa thuận / Theo chính sách công ty", indent_left=1, space_after=2)
add_body("☐  Chế độ khác: ...........................................................................................................", indent_left=1, space_after=5)

# ══════════════════════════════════════════════════════════════
# ĐIỀU 7
# ══════════════════════════════════════════════════════════════
add_heading_styled("ĐIỀU 7: CHẤM DỨT HỢP ĐỒNG THỬ VIỆC", size=13)

add_body("7.1. Trong thời gian thử việc:", bold=True, indent_left=0.5, space_after=3)
add_body("Theo Điều 27 Bộ luật Lao động 2019, trong thời gian thử việc, mỗi bên có quyền hủy bỏ Hợp đồng thử việc mà không cần báo trước và không phải bồi thường.", indent_left=0.5, space_after=5)

add_body("7.2. Khi kết thúc thời gian thử việc:", bold=True, indent_left=0.5, space_after=3)
add_body("a)  Trường hợp đạt yêu cầu: Bên A thông báo kết quả và hai bên tiến hành ký kết Hợp đồng lao động chính thức. Thời gian thử việc được tính vào thời gian làm việc tại Công ty.", indent_left=1, space_after=2)
add_body("b)  Trường hợp không đạt yêu cầu: Bên A thông báo kết quả và Hợp đồng thử việc đương nhiên chấm dứt. Bên A thanh toán đầy đủ tiền lương và các chế độ cho Bên B đến ngày chấm dứt.", indent_left=1, space_after=5)

add_body("7.3. Thời hạn thông báo kết quả:", bold=True, indent_left=0.5, space_after=3)
add_body("Bên A thông báo kết quả đánh giá cho Bên B trong vòng 03 (ba) ngày làm việc kể từ ngày kết thúc thời gian thử việc.", indent_left=0.5, space_after=5)

# ══════════════════════════════════════════════════════════════
# ĐIỀU 8
# ══════════════════════════════════════════════════════════════
add_heading_styled("ĐIỀU 8: BẢO MẬT THÔNG TIN", size=13)

add_body("8.1.  Bên B cam kết bảo mật mọi thông tin liên quan đến hoạt động kinh doanh, chiến lược, khách hàng, dữ liệu, tài liệu nội bộ và bất kỳ thông tin mật nào khác của Bên A mà Bên B tiếp cận được trong quá trình làm việc.", indent_left=0.5, space_after=3)

add_body("8.2.  Nghĩa vụ bảo mật có hiệu lực trong suốt thời gian thử việc và tiếp tục có hiệu lực trong vòng ......... (.........) năm kể từ ngày chấm dứt Hợp đồng thử việc hoặc Hợp đồng lao động.", indent_left=0.5, space_after=3)

add_body("8.3.  Bên B không được sử dụng thông tin mật của Bên A cho mục đích cá nhân hoặc cung cấp cho bên thứ ba dưới bất kỳ hình thức nào.", indent_left=0.5, space_after=3)

add_body("8.4.  Vi phạm điều khoản bảo mật, Bên B phải bồi thường thiệt hại cho Bên A theo quy định pháp luật.", indent_left=0.5, space_after=5)

# ══════════════════════════════════════════════════════════════
# ĐIỀU 9
# ══════════════════════════════════════════════════════════════
add_heading_styled("ĐIỀU 9: SỞ HỮU TRÍ TUỆ", size=13)

add_body("9.1.  Mọi sản phẩm trí tuệ, bao gồm nhưng không giới hạn: nội dung sáng tạo (bài viết, hình ảnh, video, chiến lược nội dung), tài liệu, thiết kế, ý tưởng mà Bên B tạo ra trong phạm vi công việc và sử dụng nguồn lực của Bên A đều thuộc quyền sở hữu của Bên A.", indent_left=0.5, space_after=3)

add_body("9.2.  Bên B không được sử dụng, sao chép, phát tán hoặc công bố các sản phẩm trí tuệ nêu trên cho mục đích cá nhân hoặc bên thứ ba khi chưa có sự đồng ý bằng văn bản của Bên A.", indent_left=0.5, space_after=5)

# ══════════════════════════════════════════════════════════════
# ĐIỀU 10
# ══════════════════════════════════════════════════════════════
add_heading_styled("ĐIỀU 10: ĐIỀU KHOẢN CHUNG", size=13)

add_body("10.1.  Hợp đồng này có hiệu lực kể từ ngày ký.", indent_left=0.5, space_after=3)

add_body("10.2.  Hợp đồng được lập thành 02 (hai) bản có giá trị pháp lý như nhau, mỗi bên giữ 01 (một) bản.", indent_left=0.5, space_after=3)

add_body("10.3.  Trong quá trình thực hiện, nếu phát sinh vấn đề, hai bên sẽ cùng thương lượng giải quyết trên tinh thần hợp tác. Trường hợp không thỏa thuận được, tranh chấp sẽ được giải quyết theo quy định pháp luật lao động Việt Nam.", indent_left=0.5, space_after=3)

add_body("10.4.  Những vấn đề không được quy định trong Hợp đồng này sẽ được thực hiện theo quy định của Bộ luật Lao động, nội quy lao động của Công ty và các quy định pháp luật có liên quan.", indent_left=0.5, space_after=15)

# ══════════════════════════════════════════════════════════════
# SIGNATURE BLOCK
# ══════════════════════════════════════════════════════════════

# Signature table (no borders)
sig_table = doc.add_table(rows=7, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Remove all borders
tbl = sig_table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '</w:tblBorders>'
)
tblPr.append(borders)

# Row 0: Headers
set_cell_text(sig_table.cell(0, 0), "BÊN A", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(sig_table.cell(0, 1), "BÊN B", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

set_cell_text(sig_table.cell(1, 0), "(Người sử dụng lao động)", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(sig_table.cell(1, 1), "(Người lao động)", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

set_cell_text(sig_table.cell(2, 0), "(Ký tên, đóng dấu)", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(sig_table.cell(2, 1), "(Ký tên)", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

# Rows 3-5: blank space for signatures
for row_idx in range(3, 6):
    set_cell_text(sig_table.cell(row_idx, 0), "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(sig_table.cell(row_idx, 1), "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

# Row 6: Name lines
set_cell_text(sig_table.cell(6, 0), ".........................................", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(sig_table.cell(6, 1), ".........................................", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

# Set column widths
for row in sig_table.rows:
    row.cells[0].width = Cm(7.5)
    row.cells[1].width = Cm(7.5)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), "hop-dong-thu-viec-content-strategist.docx")
doc.save(output_path)
print(f"Saved: {output_path}")
